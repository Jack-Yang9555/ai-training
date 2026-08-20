from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "实训资料" / "产品与交付" / "AI通识培训与教学应用平台技术参数及功能要求.docx"

BLUE = "1F4E79"
BLUE_DARK = RGBColor(31, 78, 121)
TABLE_HEADER = "D9D9D9"
TOTAL_FILL = "EDEDED"
VERY_LIGHT = "F7F7F7"
WHITE = "FFFFFF"
BLUE_SOFT = VERY_LIGHT
CYAN = BLUE
CYAN_SOFT = VERY_LIGHT
VIOLET = BLUE
VIOLET_SOFT = VERY_LIGHT
ORANGE = BLUE
ORANGE_SOFT = VERY_LIGHT
AMBER = BLUE
AMBER_SOFT = VERY_LIGHT
GRAY_SOFT = VERY_LIGHT
BLUE_PALE = WHITE
GOLD = BLUE
LINE = "BFBFBF"
LINE_DARK = "A6A6A6"
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=75, start=70, bottom=75, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=4, outer_color=LINE_DARK, outer_size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:color"), outer_color if edge in ("top", "left", "bottom", "right") else color)
        node.set(qn("w:sz"), str(outer_size if edge in ("top", "left", "bottom", "right") else size))


def set_paragraph_border(paragraph, edge="bottom", color=BLUE, size=8, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:color"), color)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    page_num = sect_pr.find(qn("w:pgNumType"))
    if page_num is None:
        page_num = OxmlElement("w:pgNumType")
        sect_pr.append(page_num)
    page_num.set(qn("w:start"), str(start))


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def add_page_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("AI 通识培训与教学应用平台技术参数及功能要求")
    set_run(r, 7.8, False, MUTED)
    set_p(hp, 0, 3, 1.0)
    set_paragraph_border(hp, edge="bottom", color=LINE_DARK, size=4, space=4)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("—  第 ")
    set_run(r, 8, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    r = fp.add_run(" 页  —")
    set_run(r, 8, False, MUTED)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_run(run, size=9, bold=False, color=INK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_p(p, before=0, after=2, line=1.25):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def text(doc, value="", size=9, bold=False, color=INK, before=0, after=3, align=None, line=1.35):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(value)
    set_run(r, size, bold, color)
    set_p(p, before, after, line)
    return p


def heading(doc, value, level=1):
    size = 18 if level == 1 else 11.5
    p = doc.add_paragraph()
    r = p.add_run(value)
    set_run(r, size, True, BLUE_DARK if level == 1 else INK)
    set_p(p, 3 if level == 1 else 6, 8 if level == 1 else 4, 1.15)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0.12)
        set_paragraph_border(p, edge="bottom", color=LINE_DARK, size=4, space=6)
    else:
        p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    return p


def bullets(doc, items, size=8.7):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, size, False, INK)
        set_p(p, 0, 1, 1.25)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    set_p(p, 0, 0, 1)


def callout(doc, value, fill=BLUE_SOFT, border=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    c = t.cell(0, 0)
    set_cell_width(c, 17.2)
    set_cell_shading(c, VERY_LIGHT)
    set_cell_margins(c, 130, 150, 130, 150)
    tc_pr = c._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("left", "top", "right", "bottom"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), LINE)
        borders.append(node)
    tc_pr.append(borders)
    p = c.paragraphs[0]
    r = p.add_run(value)
    set_run(r, 8.7, False, INK)
    set_p(p, 0, 0, 1.3)
    text(doc, "", after=1)


def add_table(doc, headers, rows, widths, font_size=7.4, total_rows=None):
    total_rows = set(total_rows or [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    for i, h in enumerate(headers):
        c = header.cells[i]
        set_cell_width(c, widths[i])
        set_cell_shading(c, TABLE_HEADER)
        set_cell_margins(c, 90, 65, 90, 65)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run(r, font_size, True, INK)
        set_p(p, 0, 0, 1.05)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_split(row)
        is_total = row_idx in total_rows
        for i, value in enumerate(values):
            c = row.cells[i]
            set_cell_width(c, widths[i])
            set_cell_margins(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_total:
                set_cell_shading(c, TOTAL_FILL)
            else:
                set_cell_shading(c, WHITE)
            p = c.paragraphs[0]
            r = p.add_run(str(value))
            color = BLUE_DARK if i == 1 and str(value) == "★" else INK
            set_run(r, font_size, is_total or (i == 1 and str(value) == "★"), color)
            set_p(p, 0, 0, 1.15)
            if i in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text(doc, "", after=1)
    return table


FEATURE_IDS = {
    # 仅标识平台差异化特色；容量、安全、交付等刚性要求不因重要而标星。
    "A-01", "A-05",
    "R-07",
    "T-04", "T-05", "T-11", "T-12",
    "L-03", "L-04", "L-07", "L-13", "L-14",
    "G-01", "G-02", "G-07", "G-12",
    "P-02", "P-05", "P-10",
    "F-01", "F-04", "F-08", "F-09",
    "S-01", "S-02", "S-08",
    "M-05", "M-08",
}


def parameterize_rows(rows):
    """将所有平台功能统一转换为可响应、可验收的参数表达。"""
    normalized = []
    for row in rows:
        row_id, _legacy_level, item, requirement, acceptance = row
        feature_mark = "★" if row_id in FEATURE_IDS else ""
        normalized.append([
            row_id,
            feature_mark,
            item,
            f"响应值：支持/满足；最低参数要求：{requirement}",
            acceptance,
        ])
    return normalized


def param_table(doc, rows, font_size=7.15):
    return add_table(
        doc,
        ["编号", "特色", "参数项", "功能参数（最低要求）", "验收方式"],
        parameterize_rows(rows),
        [1.15, 0.8, 3.05, 9.7, 2.5],
        font_size=font_size,
    )


def feature_cards(doc, cards):
    table = doc.add_table(rows=(len(cards) + 1) // 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="D8E1ED", size=4, outer_color="B7C5D8", outer_size=5)
    for i, (title_value, body) in enumerate(cards):
        c = table.cell(i // 2, i % 2)
        set_cell_width(c, 8.6)
        set_cell_shading(c, WHITE)
        set_cell_margins(c, 130, 150, 130, 150)
        p = c.paragraphs[0]
        r = p.add_run(title_value + "\n")
        set_run(r, 11, True, BLUE_DARK)
        r = p.add_run(body)
        set_run(r, 8.2, False, MUTED)
        set_p(p, 0, 0, 1.3)
    text(doc, "", after=1)


def add_hyperlink(paragraph, label, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Microsoft YaHei")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = label
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
sec = doc.sections[0]
configure_page(sec)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(9.2)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.35
for style_name in ("List Bullet", "List Number"):
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(8.7)

doc.core_properties.title = "AI通识培训与教学应用平台技术参数及功能要求"
doc.core_properties.subject = "融合平台特色、48课时课程与前沿教学设计的可验收技术参数"
doc.core_properties.author = "启境 AI Learning OS 项目组"

# Cover
text(doc, "TECHNICAL SPECIFICATION", size=10.5, bold=True, color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, before=35, after=18)
text(doc, "AI 通识培训与教学应用平台", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.1)
text(doc, "技术参数及功能要求", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, line=1.1)
text(doc, "融合平台特色、前沿教学设计与 48 课时课程内容的建设参数", size=11.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
feature_cards(doc, [
    ("Learning OS", "教、学、测、评、练、研、管一体化，以真实教学任务和能力证据组织平台。"),
    ("双空间成长", "教师教学工作与个人 AI 成长相互独立又可共享成果证据。"),
    ("课程即能力", "16 个课程模块、24 次课、项目成果、评测集和补学任务进入统一闭环。"),
    ("可信 AI", "来源、权限、版本、置信度、证据充分度、人工复核和生成标识贯穿全过程。"),
])
callout(doc, "适用场景：高职院校 AI 通识课程建设、教师 AI 能力培训、校级教学运行管理、AI 教学应用实训与研究创新。本文档可作为项目立项、采购技术需求、招标参数、供应商响应和验收测试的基础文本。")
meta_rows = [
    ["文档版本", "V1.0", "编制日期", "2026 年 7 月 31 日"],
    ["建议规模", "100 名教师账号", "峰值并发", "80 人同时调用 AI"],
    ["课程基线", "48 课时 / 16 模块 / 24 次课", "内容留存", "不少于 3 年"],
    ["部署方式", "校内部署 + 云端模型", "服务模式", "7×8 运维"],
]
add_table(doc, ["项目", "参数", "项目", "参数"], meta_rows, [2.3, 6.3, 2.3, 6.3], font_size=8.3)

# 1
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
configure_page(body_section)
add_page_header_footer(body_section)
set_page_number_start(body_section, 1)
heading(doc, "一、文档定位与参数使用说明")
callout(doc, "核心定位：平台不是工具集合，而是围绕“课程知识—教学任务—个性化学习—多元测评—诊断报告—补学回流—教师成长—学校治理”形成连续数据闭环的 AI Learning OS。")
heading(doc, "1. 参数标识与响应口径", level=2)
add_table(doc, ["标记", "含义", "使用规则"], [
    ["★", "平台特色参数", "仅用于标识启境 AI Learning OS 的差异化特色，包括双空间、证据闭环、课程适配、可信 AI、教师成长、Harness 与 Agent Eval 等。"],
    ["空白", "平台通用参数", "属于完整平台必须具备的常规功能、性能、安全、交付或运维要求，不因重要程度重复标星。"],
    ["统一句式", "参数化响应", "全部功能采用“响应值：支持/满足；最低参数要求：……”表达，并给出对应验收方式；供应商须逐项响应，不得以产品介绍替代。"],
], [2.0, 4.0, 11.2], font_size=8.4)
heading(doc, "2. 建设原则", level=2)
feature_cards(doc, [
    ("任务对象优先", "教师先处理课程、班级、测评、报告和实训对象，不以传统后台菜单堆叠能力。"),
    ("人机责任清晰", "AI 提供草稿、证据和建议；教师负责发布、评分、关键结论与最终责任。"),
    ("证据驱动", "学习、作品、测评、助教对话、实训与研究成果进入可追溯证据台账。"),
    ("开放可替换", "模型、知识库、向量库、智能体和工具协议可替换，避免绑定单一厂商。"),
    ("安全内生", "隐私、版权、内容标识、权限、人工接管和审计不是附加模块，而是业务字段。"),
    ("持续改进", "支持课程版本、模型评测、干预回流、结业复测与 30/90 天延迟评价。"),
])
heading(doc, "3. 平台特色摘要", level=2)
bullets(doc, [
    "一套平台同时服务教师教学工作、教师个人成长、学生学习和学校教学运行管理。",
    "同一课程知识源同时驱动知识图谱、AI 助教、智能出题、组卷、批改、诊断和补学。",
    "课程内容覆盖 AI 基础、上下文工程、Agent Skill、RAG、MCP、状态与记忆、智能体、Harness、Agent Eval 与安全治理。",
    "形成“学、测、训、用、研、评”教师成长闭环和“教、学、测、评、练”学生学习闭环。",
])

# 2
page_break(doc)
heading(doc, "二、总体架构与部署参数")
architecture_rows = [
    ["A-01", "★", "总体定位", "应建设面向师生教学、教师成长与学校教学运行的 Learning OS，一级能力至少覆盖教、学、测、评、练、研、管；不得仅以独立 AI 工具导航替代教学闭环。", "现场演示"],
    ["A-02", "★", "混合部署", "支持校内部署业务系统、账号权限、数据库、对象存储、向量索引、审计和模型网关；经审批的文本、视觉、图片、视频模型可通过云端 API 弹性调用。", "架构与部署核验"],
    ["A-03", "★", "模型网关", "所有模型调用应经服务端统一网关，不得在浏览器暴露密钥；支持模型路由、限流、超时、重试、熔断、配额、成本统计和供应商切换。", "接口与故障测试"],
    ["A-04", "▲", "分层架构", "至少包含用户体验层、教学业务层、AI 能力层、数据与知识层、集成层、安全治理层和基础设施层，各层边界清晰。", "设计文档"],
    ["A-05", "★", "统一知识底座", "课程大纲、讲义、案例、习题、量规和制度规范应形成统一知识源，并保留来源、版本、适用模块、知识点、权限和更新时间。", "数据抽查"],
    ["A-06", "▲", "可替换 AI 服务", "文本模型、视觉理解、图片生成、视频生成、Embedding、Rerank、OCR、文档解析和内容安全能力应通过适配层接入。", "替换演示或接口文档"],
    ["A-07", "▲", "异步任务", "图片、视频、批量报告、知识库解析、批量批改等长任务应异步执行，支持进度、取消、失败原因、重试和完成通知。", "任务演示"],
    ["A-08", "▲", "工作流与智能体", "支持低代码或配置化编排模型、知识、Skill、工具、状态、规则、护栏和人工确认节点；运行轨迹可查看。", "场景演示"],
    ["A-09", "▲", "开放协议", "工具连接支持标准 REST/API，并预留 MCP 等开放协议；多智能体协作和 A2A 可作为扩展能力，不作为公共通识必修工程要求。", "接口核验"],
    ["A-10", "★", "数据主权", "甲方课程资料、提示词、生成内容、作品、评分、报告、日志和研究成果归甲方；支持导出、迁移、归档和按权限删除。", "合同与数据测试"],
    ["A-11", "▲", "多端适配", "支持主流桌面浏览器；移动端可完成核心查看、学习、作答、报告和任务操作；1440px 与 390px 不得出现关键遮挡或横向溢出。", "浏览器验收"],
    ["A-12", "▲", "可观测性", "提供业务日志、模型日志、任务队列、调用耗时、Tokens、失败率、低置信度、人工接管、存储和备份状态监控。", "监控演示"],
]
param_table(doc, architecture_rows)

# 3
page_break(doc)
heading(doc, "三、用户、角色与工作空间参数")
role_rows = [
    ["R-01", "★", "教师角色", "教师应具备智能教学、学习促进、测评管理、教学诊断、AI 研究、自由实训，以及“教学工作 / 我的成长”双空间。", "角色演示"],
    ["R-02", "★", "学生角色", "学生应具备课程导学、个性化学习、我的测评、成长报告和自由实训；不得看到教师批改、班级报告和研究数据。", "权限测试"],
    ["R-03", "★", "教学管理者", "管理者应使用运行总览、课程班级、学习质量、测评运行、AI 应用等独立任务面，只查看聚合数据和匿名证据。", "权限测试"],
    ["R-04", "★", "集中权限策略", "入口权限、数据权限和状态调用权限应由统一策略约束；前端隐藏不得代替服务端鉴权。", "接口越权测试"],
    ["R-05", "▲", "对象优先首页", "登录后先展示真实课程、班级、测评、报告、培训和实训对象，卡片包含状态、时间、进度、指标和唯一主操作。", "页面演示"],
    ["R-06", "▲", "上下文继承", "教师选择的班级在教、学、评之间继承，学生选择的课程在教、学、测、评之间继承；切换对象清除不兼容数据。", "交互测试"],
    ["R-07", "★", "双空间隔离", "教师教学班级上下文与教师个人学习课程、笔记、路径、实训和能力报告分别保存，切换空间不得相互覆盖。", "状态测试"],
    ["R-08", "▲", "课程隔离", "两门及以上课程的导学、任务、测评、报告、补学和知识库状态应按 courseId 或等效主键隔离。", "数据测试"],
    ["R-09", "▲", "组织范围", "支持学校、院系、专业、课程、班级、教师等组织层级，并支持全校及院系范围聚合。", "配置演示"],
    ["R-10", "▲", "匿名证据", "管理者查看个人明细时使用“教师/学生匿名编号”，不得暴露无关实名、联系方式、原始作品或个人报告。", "权限抽查"],
    ["R-11", "—", "会话与重置", "支持保存当前工作对象和未完成任务；演示或试点环境可提供受控重置，生产环境不得以重置替代持久化。", "状态测试"],
    ["R-12", "▲", "无障碍", "支持键盘操作、清晰焦点、语义标题、替代文本、字幕、合理对比度、触控尺寸和减少动画偏好。", "可访问性检查"],
]
param_table(doc, role_rows)

# 4
page_break(doc)
heading(doc, "四、教师智能教学与课程建设参数")
teaching_rows = [
    ["T-01", "★", "标准课程结构", "支持将 AI 通识课程配置为 16 个模块、24 次课、48 课时；每次课支持目标、知识点、流程、资源、活动、题目、量规和课后任务。", "课程演示"],
    ["T-02", "★", "备课任务流", "教师“教”阶段二级任务至少为：智能备课 → 教学实施 → 课程资源 → AI 助教；知识图谱归入课程资源的知识结构。", "导航演示"],
    ["T-03", "▲", "情境适配", "支持院校教学、教师培训等情境，以及对象、基础、人数、授课方式、时长、难度和重点配置。", "生成演示"],
    ["T-04", "★", "AI 适配依据", "AI 生成的教学方案必须展示标准方案与适配方案在目标、难度、节奏、活动、资源和题目方面的差异及依据。", "结果核验"],
    ["T-05", "★", "人工采纳发布", "AI 方案首先为草稿；教师可修改并采纳，发布后转为正式资源，保留 AI 原稿、人工修改和版本链。", "版本测试"],
    ["T-06", "▲", "四类备课成果", "采纳方案后至少形成情境化教案、课堂活动包、配套题目集、教学进度方案四类结构化成果。", "成果抽查"],
    ["T-07", "▲", "教学实施", "支持调整各教学步骤时长、资源版本、题目、知识点和准备度检查；支持使用标准方案直接上课。", "课堂演示"],
    ["T-08", "▲", "课程资源", "提供资源总览、标准资源、AI 生成、知识结构等视图；资源支持来源、版本、状态、权限、标签、知识点和引用关系。", "资源核验"],
    ["T-09", "★", "知识库入库", "支持大纲、讲义、案例、习题、量规、操作任务单和术语表导入；记录解析、分段、重复率、覆盖率和异常。", "导入演示"],
    ["T-10", "▲", "知识图谱", "知识库就绪后方可生成知识图谱；节点和关系支持人工核验、修订、版本保存和知识点关联。", "状态测试"],
    ["T-11", "★", "AI 助教搭建", "至少包含知识绑定、教学人格、数字人/形态、测试发布四步；知识库为必选底座，知识图谱为可选增强。", "配置演示"],
    ["T-12", "★", "有引用回答", "AI 助教回答应展示课程来源、章节、知识点、置信度；超出范围时明确知识边界并建议教师确认。", "问答测试"],
    ["T-13", "▲", "助教发布", "支持测试集、版本、开放范围、课程入口、链接或二维码；发布前检查引用、边界、内容安全和人工接管。", "发布演示"],
    ["T-14", "▲", "助教运营", "展示解决率、知识库命中率、高频问题、低置信度回答、越界问题、人工接管和问题趋势。", "运营报表"],
    ["T-15", "▲", "多模态资源", "支持图文、流程图、海报、音频、短视频或数字人脚本生成；必须保留提示词、模型、版本、人工修改、版权与无障碍检查。", "成果抽查"],
    ["T-16", "▲", "课程版本", "支持课程、模块、课次和资源版本；发布版本与编辑草稿隔离，历史引用可回溯。", "版本测试"],
    ["T-17", "—", "课程复用", "支持课程复制、模板复用和校本化调整，复制后生成独立标识，避免状态串扰。", "操作演示"],
    ["T-18", "▲", "证据入口", "备课方案、AI 助教配置、资源生成和发布结果均可打开统一证据详情。", "证据抽查"],
]
param_table(doc, teaching_rows)

# 5
page_break(doc)
heading(doc, "五、学习促进、测评与诊断参数")
learning_rows = [
    ["L-01", "★", "学习促进任务流", "教师“学”阶段至少为：班级态势 → 掌握诊断 → 分层干预；不得复用学生个人学习操作界面。", "导航与权限测试"],
    ["L-02", "▲", "多周期观察", "支持本课、本周、本单元等观察周期，默认使用过程指标，不以正式成绩和排名替代学习过程。", "筛选演示"],
    ["L-03", "★", "多源学习证据", "按目标聚合学习任务、课堂练习、作品迭代、测评过程和 AI 助教对话，显示来源、充分度和更新时间。", "证据抽查"],
    ["L-04", "★", "诊断区分", "明确区分“尚未掌握”和“证据不足”；低置信度或证据不足结果不得直接触发高影响自动决策。", "规则测试"],
    ["L-05", "▲", "动态分层", "支持补强、巩固、拓展三类动态分组；分组依据可解释、接收对象可调整，避免公开贴标签。", "分组演示"],
    ["L-06", "★", "干预发布", "AI 生成分层任务草稿，教师可调整目标、资源、题目、难度、时长、截止日期和接收对象后发布。", "闭环演示"],
    ["L-07", "★", "学习回流", "学生完成分层任务或补学后，应回流班级态势、掌握诊断和下一轮教学决策。", "跨角色测试"],
    ["L-08", "▲", "个性化学习", "学生支持随堂学/课后学、学习路径、作品版本、知识点难度、练习、收藏和笔记。", "学习演示"],
    ["L-09", "▲", "五级难度", "练习难度支持五级，并解释认知层级、前置知识、历史表现和任务复杂度。", "规则核验"],
    ["L-10", "★", "多场景测评", "支持随堂测、单元测、正式考试、实训和综合项目；题目与课程目标、知识点和量规关联。", "测评演示"],
    ["L-11", "▲", "七类题型", "支持单选、多选、判断、填空、简答、案例分析和创作实操题；主观题支持证据和量规。", "题库演示"],
    ["L-12", "★", "组卷蓝图", "支持按总分、知识覆盖、难度曲线、题型、认知层级和评分主体组卷，并检查相似题、量规和总分。", "组卷测试"],
    ["L-13", "★", "多主体评分", "综合项目可采用机器客观评分、AI 作品评分、教师人工评分；默认比例 30/50/20，并允许课程级配置。", "评分演示"],
    ["L-14", "★", "评分证据", "AI 作品评分至少展示任务符合度、专业准确性初检、方法与工具、迭代证据、内容安全及量规条目。", "证据抽查"],
    ["L-15", "★", "人工复核", "教师可调整 AI 分和人工分，最终分实时重算；AI 原分、教师修改理由和最终分同时保留。", "评分测试"],
    ["L-16", "▲", "报告体系", "教师报告包括提交、成绩/达成、复核、难度、分布、知识掌握和共性误区；学生报告只展示个人证据与建议。", "报告演示"],
    ["L-17", "★", "补学回流", "报告建议至少分立即补强、继续巩固、拓展挑战；生成补学任务后回到课后学习并可完成回流。", "闭环测试"],
    ["L-18", "▲", "过程性评价", "平台支持学习过程、知识测评、核心实训作品和综合项目等多元评价，不以点击次数代替能力。", "评价方案核验"],
]
param_table(doc, learning_rows)

# 6
page_break(doc)
heading(doc, "六、教师成长、实训与研究参数")
growth_rows = [
    ["G-01", "★", "成长闭环", "教师成长空间形成：摸底测评 → 详细能力报告 → 个性化培训方案 → 课程/自主学习 → AI 实训 → 研究/工具开发 → 成长报告与复测。", "全流程演示"],
    ["G-02", "★", "5+3 能力模型", "五个计分维度至少为 AI 基础认知、提示词与多模态、知识库与智能体、教学融合、研究创新；设置事实核验、数据版权、人工责任等通关门槛。", "报告核验"],
    ["G-03", "▲", "能力等级", "支持 L1 入门者、L2 实践者、L3 创新者、L4 引领者或等效等级，等级规则可配置并有证据支撑。", "规则测试"],
    ["G-04", "★", "真实摸底", "测评包含理论、情境和实操；客观题真实作答，实操题提交提示词、截图、文件、测试记录或人工确认。", "作答演示"],
    ["G-05", "▲", "个性化计划", "基于摸底报告从已导入课程生成培训计划，支持调整学习顺序、每周投入和选修内容。", "计划演示"],
    ["G-06", "▲", "自主学习", "支持 AI 推荐路径与自主路径并存；记录课程、资源进度、收藏、笔记、继续学习和完成证据。", "学习演示"],
    ["G-07", "★", "成果证据台账", "课程学习、自由实训、教学实施、评分复核、分层干预、研究成果和复测等至少七类成果进入个人台账并可下钻。", "证据抽查"],
    ["G-08", "▲", "18 项真实任务", "支持将备课资源、命题评价、知识库智能体、教学研究等四条任务链组织为不少于 18 项真实工作任务。", "任务地图"],
    ["G-09", "▲", "11 项平台实训", "自由实训任务链至少覆盖创建知识库、检索优化、生成图谱、配置智能体/数字人、发布助教、出题、实训、组卷、批改和学情分析。", "任务演示"],
    ["G-10", "▲", "任务前置关系", "每项任务显示频率、难度、标准时长、前置任务、输入、操作、产出、量规和完成状态；按前置关系解锁。", "状态测试"],
    ["G-11", "▲", "实训成果回流", "自由实训结果进入教师实训记录和能力报告，记录工具、任务、用时、得分、作品、版本和人工确认。", "回流演示"],
    ["G-12", "★", "AI 研究工作台", "仅教师可见，至少包含研究项目、共享 AI 工具、科研智能体、成果与伦理；学生和管理者不得访问研究数据。", "权限测试"],
    ["G-13", "▲", "研究四阶段", "支持问题与设计、文献与证据、分析与干预、成果与复核四阶段；阶段产出、来源、参数和人工判断可追溯。", "研究演示"],
    ["G-14", "▲", "科研智能体", "支持研究项目、已导入资料、角色、工具、测试集、版本、成果关联、匿名范围和 AI 使用披露。", "配置演示"],
    ["G-15", "★", "研究伦理", "成果至少检查数据匿名、引用回溯、AI 使用披露、模型/提示参数归档和人工复核。", "合规抽查"],
    ["G-16", "▲", "延迟复测", "支持首次摸底、过程评价、结业复测及可选 30/90 天延迟复测，展示能力变化和新任务迁移证据。", "报告演示"],
]
param_table(doc, growth_rows)

# 7
page_break(doc)
heading(doc, "七、48 课时 AI 通识课程内容与平台映射参数")
text(doc, "平台应将课程内容直接配置为可教学、可实训、可评价、可沉淀的运行单元，而不是仅提供静态课程目录。每个模块应包含目标、知识点、资源、任务、工具、成果、量规、风险边界和补学建议。", size=9.2, color=MUTED)
course_rows = [
    ["C-01", "AI 基础、职业变革与人机分工", "4", "课程导学、课前诊断、学习路径、深度研究/模型评测", "AI 应用辨识图、岗位任务人机分工表"],
    ["C-02", "数据、机器学习与模型评价", "4", "随堂学、研究数据分析室、模型评测场、单元测", "数据质量诊断表、模型选择说明"],
    ["C-03", "生成式 AI 原理与可信使用", "4", "按需资源、模型评测、风险情境测评、成长报告", "AI 输出核验记录、可信使用清单"],
    ["C-04", "上下文工程", "2", "提示词实验室、AI 助教、作品版本", "岗位任务上下文包 v1/v2"],
    ["C-05", "提示词与结构化输出", "4", "提示词实验室、数据分析、模型评测", "三轮提示词记录、结构化输出模板"],
    ["C-06", "Agent Skill 设计", "2", "提示词实验室、Vibe Coding、成果模板", "入门级 SKILL.md、试用反馈"],
    ["C-07", "多模态内容生成", "2", "多模态内容工坊、识别、作品版本", "多模态作品、版权和无障碍检查"],
    ["C-08", "语音与数字人", "2", "文生视频、多模态工坊、数字人形态", "语音或数字人方案、字幕与披露记录"],
    ["C-09", "计算机使用智能体", "2", "仿真页面、智能体搭建、工具连接", "受控操作轨迹、人工确认点清单"],
    ["C-10", "课程知识库与 RAG", "4", "课程资源、知识库 RAG、AI 助教、单元实训", "来源清单、6–9 题检索测试报告"],
    ["C-11", "工具调用与 MCP", "2", "MCP 工具连接台、Dify、权限矩阵", "工具说明卡、输入输出样例、权限表"],
    ["C-12", "状态、记忆与任务恢复", "2", "智能体搭建、Dify、运行轨迹", "状态流转图、断点恢复和幂等测试"],
    ["C-13", "智能体设计与低代码开发", "4", "智能体搭建、Dify、MCP、护栏", "智能体原型、权限风险矩阵、人机交接"],
    ["C-14", "Harness 与运行可观测性", "2", "Vibe Coding、模型评测、轨迹与日志", "Harness 画布、运行轨迹标注"],
    ["C-15", "智能体评测与安全", "4", "模型评测、MCP、红队测试、评分证据", "不少于 8 例评测集、红队报告、原型 v2"],
    ["C-16", "专业智能体综合项目", "4", "作品版本、模型评测、成长报告、答辩", "AI 职业应用成果包、验收和治理清单"],
    ["", "合计", "48", "16 个模块 / 24 次课", "完整 AI 职业应用成果包"],
]
add_table(doc, ["编号", "模块主题", "课时", "平台能力映射", "核心成果"], course_rows, [1.1, 4.0, 1.0, 6.0, 5.1], font_size=7.0, total_rows={16})

# 8
page_break(doc)
heading(doc, "八、前沿教学设计参数")
pedagogy_rows = [
    ["P-01", "★", "人本与教师主体", "平台应强化人的能动性和责任，不以 AI 替代教师教学判断；高影响发布、评分、研究结论和不可逆操作必须保留人工确认。", "流程测试"],
    ["P-02", "★", "教师-AI-学生关系", "支持教师设计、AI 辅助、学生实践、教师复核的三方互动；AI 助教不直接替学生完成计分任务。", "课堂场景"],
    ["P-03", "▲", "Acquire–Deepen–Create", "课程与教师成长路径应覆盖获取基础、深化应用、创造创新三个进阶层次，并映射到平台能力等级和任务证据。", "课程与报告核验"],
    ["P-04", "★", "任务驱动", "每个教学模块以真实或仿真的职业任务组织内容，包含输入、步骤、工具、成果、检查点、量规和风险边界。", "课程抽查"],
    ["P-05", "★", "项目贯穿", "综合项目从智能体选题阶段启动，持续补充权限、Harness、评测和治理证据，不得压缩为结课一次性作业。", "项目轨迹"],
    ["P-06", "▲", "做中学", "短讲解、正反案例、教师示范、个人/小组实践、展示互评、即时检测和证据归档形成每次课基本节奏。", "教案抽查"],
    ["P-07", "▲", "证据中心设计", "学习目标、任务、作品、测评、对话、实训、研究和人工复核形成证据链，报告结论可以回到原始证据。", "报告下钻"],
    ["P-08", "▲", "掌握学习", "基于目标达成与证据充分度提供补强、巩固、拓展路径；补学完成后再次采集证据并更新诊断。", "闭环测试"],
    ["P-09", "▲", "自适应与可解释", "个性化推荐应说明薄弱点、前置知识、历史表现、任务复杂度和推荐原因，允许教师与学习者调整。", "推荐核验"],
    ["P-10", "★", "真实性评价", "采用过程证据、任务作品、知识测评、综合项目和答辩；AI 评分只作为辅助，争议与高风险结果人工复核。", "评价方案"],
    ["P-11", "▲", "作品集评价", "保留提示词、上下文、Skill、配置、版本、评测、失败案例、人工修改和反思，形成可复核作品集。", "作品抽查"],
    ["P-12", "★", "安全通关门槛", "敏感信息、未披露 AI 使用、未核验关键事实、高风险自动决策、越权工具、失败循环和侵权内容应触发暂不通过。", "规则测试"],
    ["P-13", "▲", "跨专业迁移", "共同知识、评价标准不变，任务输入支持制造、计算机、医药、经管、艺术、农林等专业情境替换。", "案例配置"],
    ["P-14", "▲", "分层而不贴标签", "分层依据仅用于支持策略，不公开年龄、专业或能力标签；学生个人报告不展示班级排名和同伴比较。", "页面与权限核验"],
    ["P-15", "▲", "通用设计与可及性", "为视觉、听觉、阅读或操作困难学习者提供替代文本、字幕、键盘操作、延长时间和等效任务。", "可访问性测试"],
    ["P-16", "▲", "学习分析改进", "课程团队可基于诊断前后变化、作品人工修改质量、引用命中、工具成功、接管和安全通关等指标迭代课程。", "运营报告"],
]
param_table(doc, pedagogy_rows)

# 9
page_break(doc)
heading(doc, "九、前沿 AI 能力与实训工具参数")
frontier_rows = [
    ["F-01", "★", "上下文工程", "支持目标、背景、资料、术语、示例、限制、状态、质量标准和版本组织；提供上下文缺失、冲突、过时、过载和污染检查。", "实训演示"],
    ["F-02", "▲", "结构化输出", "支持表格、JSON 或 Schema 约束，校验字段、类型、必填、异常值和失败回退；确定性计算优先调用工具。", "校验测试"],
    ["F-03", "▲", "Agent Skill", "支持按 SKILL.md、references、scripts、assets 等结构组织可复用技能，记录触发条件、输入、步骤、边界、依赖和验收标准。", "成果核验"],
    ["F-04", "★", "RAG 检索", "支持文档切分、向量化、关键词/混合检索、重排、引用和知识边界；提供正常、模糊、无答案和版本冲突测试。", "检索测试"],
    ["F-05", "▲", "智能体定义", "智能体至少包含目标、指令、上下文、知识、Skill、工具、状态、规划、权限、护栏和人工接管。", "配置核验"],
    ["F-06", "▲", "MCP 教学", "支持理解资源、提示和工具等基本概念，并在仿真或受控环境配置只读工具；不要求所有学习者开发 MCP 服务。", "实训演示"],
    ["F-07", "▲", "状态与记忆", "区分当前上下文、会话记忆、业务状态和长期知识；支持状态流转、断点保存、超时重试、幂等和失败转人工。", "故障测试"],
    ["F-08", "★", "Harness", "将运行环境、上下文、知识、工具、规则、日志、反馈和恢复机制组织为可查看画布或配置，并可关联运行轨迹。", "轨迹核验"],
    ["F-09", "★", "Agent Eval", "支持固定评测集重复执行，记录成功率、引用、耗时、成本、工具调用、人工接管和失败类别；至少覆盖正常、模糊、越界、攻击、高风险和故障。", "评测演示"],
    ["F-10", "★", "红队测试", "支持提示词注入、数据泄露、越权调用、错误工具结果、失败循环和高风险请求等测试，修订护栏后复测。", "安全实训"],
    ["F-11", "▲", "计算机使用智能体", "仅在隔离仿真环境执行观察、规划、填写、检查和确认；登录、验证码、敏感字段、最终提交和不可逆操作转人工。", "沙箱演示"],
    ["F-12", "▲", "多模态", "支持文本、图片、音频、视频的生成与理解任务；成果包含提示词、来源、版权、字幕、替代文本、人工核验和 AI 标识。", "成果抽查"],
    ["F-13", "▲", "模型评测", "支持同一任务、同一输入和同一量规对比多模型的质量、稳定性、速度、成本、隐私与安全，保留选择依据。", "评测报告"],
    ["F-14", "—", "多智能体与 A2A", "作为专业拓展支持智能体能力卡、任务委派、通信、失败回退和人工监督；不得作为全体通识学习者的强制编码要求。", "拓展示例"],
    ["F-15", "▲", "深度研究", "支持研究问题、检索范围、证据矩阵、来源核验、冲突证据、过程记录和人工结论，禁止伪造引用。", "研究演示"],
    ["F-16", "▲", "Vibe Coding", "支持基于自然语言生成轻量网页或工具原型，同时要求需求、验收标准、依赖、测试、修改记录和安全检查。", "原型验收"],
]
param_table(doc, frontier_rows)
page_break(doc)
heading(doc, "共享 AI 工具目录", level=2)
tool_rows = [
    ["1", "深度研究助手", "问题拆解、检索、证据矩阵和研究简报"],
    ["2", "文献研读与证据库", "来源整理、引用回溯、证据冲突与摘录"],
    ["3", "研究数据分析室", "结构化分析、确定性计算和可视化说明"],
    ["4", "提示词实验室", "上下文、提示、结构化输出和 A/B 测试"],
    ["5", "多模态内容工坊", "图文、音频、课件素材和无障碍检查"],
    ["6", "文生视频工坊", "短视频、分镜、字幕、授权与 AI 披露"],
    ["7", "知识库 RAG", "资料入库、检索测试、引用和边界"],
    ["8", "Dify 应用实验台", "低代码工作流、知识、工具和节点编排"],
    ["9", "智能体搭建台", "角色、状态、工具、护栏和人工接管"],
    ["10", "MCP 工具连接台", "工具说明、权限、输入输出和异常测试"],
    ["11", "多模态识别", "图片、文档、界面和视频内容理解"],
    ["12", "Vibe Coding 工作台", "轻量网页/工具原型与测试修改记录"],
    ["13", "模型评测场", "模型效果、稳定、速度、成本和安全对比"],
]
tool_param_rows = [
    [idx, name, f"响应值：支持/满足；最低参数要求：{capability}。", "现场演示"]
    for idx, name, capability in tool_rows
]
add_table(
    doc,
    ["序号", "工具", "功能参数（最低要求）", "验收方式"],
    tool_param_rows,
    [1.0, 3.5, 10.0, 2.7],
    font_size=7.7,
)

# 10
page_break(doc)
heading(doc, "十、可信 AI、数据安全与合规参数")
governance_rows = [
    ["S-01", "★", "统一证据字段", "关键 AI 结果至少记录来源、知识点、输入摘要、模型/服务版本、提示版本、量规、结论、置信度、证据充分度、权限范围、匿名状态、人工复核、修改记录和时间。", "字段抽查"],
    ["S-02", "★", "人工复核状态", "支持待复核、已确认、已修改、已驳回等状态；AI 原值与最终值同时保留。", "状态测试"],
    ["S-03", "★", "低置信度治理", "低置信度、无引用、证据冲突或证据不足应显式标记并进入人工处理，不得静默输出确定性结论。", "规则测试"],
    ["S-04", "★", "最小必要传输", "向云端模型只发送完成任务所需最小数据；可配置脱敏、敏感词/实体检测、禁止字段和传输审计。", "接口抓包与日志"],
    ["S-05", "★", "密钥安全", "模型、存储、数据库和第三方工具密钥由服务端密钥管理，支持轮换、权限隔离和泄露应急；不得写入前端代码。", "安全检查"],
    ["S-06", "★", "个人信息保护", "支持个人信息查阅、更正、导出和删除流程；对真实学生、教师和研究数据实行分级分类与最小权限。", "流程测试"],
    ["S-07", "▲", "版权与授权", "资源和生成内容记录来源、授权、适用范围、AI 参与方式、人工修改和发布确认；高风险或不明版权内容阻止公开发布。", "内容抽查"],
    ["S-08", "★", "生成内容标识", "对导出的 AI 生成/合成文本、图片、音频、视频按适用法规和标准提供显式提示、元数据或等效标识，并防止恶意删除。", "导出核验"],
    ["S-09", "★", "内容安全", "支持输入和输出安全检测、违法违规内容处置、用户警示、功能限制、事件记录和申诉/举报流程。", "安全测试"],
    ["S-10", "▲", "高风险边界", "医疗、工程、财务、学生处分等高风险场景明确提示正式规范和人工确认，AI 不直接作最终决定。", "情境测试"],
    ["S-11", "▲", "工具权限", "工具按只读/写入、可逆/不可逆、低/中/高风险分级；写入、发送、支付、删除和外部发布必须人工确认。", "权限测试"],
    ["S-12", "★", "审计日志", "记录登录、权限、课程发布、模型调用、工具调用、评分修改、数据导出、删除、配置变化和异常事件；日志防篡改且可检索。", "审计抽查"],
    ["S-13", "▲", "匿名研究", "研究与管理数据默认匿名/去标识；成果导出前检查再识别风险、来源授权和 AI 使用披露。", "研究抽查"],
    ["S-14", "▲", "模型与供应商管理", "记录模型备案/合规信息、供应商、版本、适用场景、价格、数据策略、可用区和替代方案。", "台账核验"],
    ["S-15", "▲", "风险闭环", "预警支持待处理、已交办、已解决等状态，并可从管理关注回流教师干预、学习者完成和指标更新。", "跨角色演示"],
    ["S-16", "▲", "AI 事件管理", "提供模型异常、错误引用、越权、敏感信息、内容标识缺失和供应商不可用等事件的登记、处置、复盘和改进。", "应急演练"],
]
param_table(doc, governance_rows)

# 11
page_break(doc)
heading(doc, "十一、学校教学运行管理与质量参数")
management_rows = [
    ["M-01", "★", "运行总览", "展示课程、班级、教师、任务、测评、报告、AI 应用、预警和服务健康的聚合指标，支持组织范围筛选。", "管理演示"],
    ["M-02", "▲", "课程班级", "支持课程开设、班级、教师、资源就绪、学习进度、测评状态和异常对象统计。", "报表核验"],
    ["M-03", "▲", "学习质量", "支持目标达成、证据充分度、补学完成、作品迭代、知识点趋势和不同群体差异，不展示无授权实名。", "报表核验"],
    ["M-04", "▲", "测评运行", "支持考试/单元测状态、提交、批改、人工复核、争议和异常任务统计。", "报表核验"],
    ["M-05", "★", "AI 应用治理", "展示 Tokens、费用、模型、成功率、响应、引用命中、低置信度、人工接管、内容安全和工具调用指标。", "运营演示"],
    ["M-06", "▲", "教师成长管理", "仅展示授权范围内的培训参与、任务完成、能力分布、成果类型和支持需求；不得查看个人答案、笔记和研究私密内容。", "权限测试"],
    ["M-07", "▲", "成果资产", "支持课程模板、提示词、量规、知识库、Skill、智能体、评测集和研究成果的校级归档、审核、复用和版本管理。", "资产演示"],
    ["M-08", "★", "预警干预", "预警包含对象范围、证据、责任人、截止时间、状态和干预结果；支持交办、跟踪、解决和指标回流。", "闭环演示"],
    ["M-09", "▲", "质量指标", "至少支持摸底/结业变化、任务达成、人工修改质量、引用命中、工具成功、人工接管、安全通关和满意度。", "指标核验"],
    ["M-10", "▲", "年度复盘", "支持按学期/年度导出课程运行、教师发展、AI 成本、风险事件、成果资产和改进计划。", "导出测试"],
    ["M-11", "—", "多组织比较", "允许在匿名聚合前提下比较院系、专业、课程和周期；小样本数据应隐藏或合并。", "隐私测试"],
    ["M-12", "▲", "管理权限", "管理者不得进入教案编辑、评分修改、学生原始作答、教师研究配置和个人成长内容。", "越权测试"],
]
param_table(doc, management_rows)

# 12
page_break(doc)
heading(doc, "十二、性能、容量、存储与运维参数")
technical_rows = [
    ["K-01", "★", "账号容量", "首期支持不少于 100 名教师正式账号，并具备横向扩展到更多教师和学生账号的架构能力。", "账号测试"],
    ["K-02", "★", "AI 峰值并发", "支持不少于 80 名教师同时发起 AI 任务；应用网关按不低于 100 路在途请求能力设计。", "压力测试"],
    ["K-03", "▲", "文本配额", "建议企业文本模型配额达到 150–200 万 TPM，可通过队列和路由控制智能体扇出；实际值写入合同附件。", "配额与压测"],
    ["K-04", "★", "业务响应", "校内网络正常条件下，普通业务页面 95% 请求响应时间不高于 3 秒。", "性能报告"],
    ["K-05", "▲", "AI 响应", "文本任务支持流式返回；上游模型完整生成时间可监控、超时、取消和重试，慢任务进入可见队列。", "接口测试"],
    ["K-06", "▲", "模型容量", "三年标准测算约 47.25 亿 Tokens，建议采购容量不低于约 65 亿等价 Tokens，并设置 30% 系统提示、重试、工具链和增长余量。", "资源台账"],
    ["K-07", "▲", "多媒体容量", "三年控制目标可按不少于 10 万次图片、8,000 次短视频或等价预算池设计，支持按组织和账号限额。", "配额测试"],
    ["K-08", "★", "内容留存", "课程资料、生成内容、作品、评分、报告和关键审计数据保存不少于 3 年；留存期可配置。", "数据抽查"],
    ["K-09", "▲", "对象存储", "建议提供不少于 10TB 可用对象存储；支持版本、校验、生命周期、归档和按权限下载。", "容量核验"],
    ["K-10", "▲", "备份容量", "建议独立备份可用容量不少于 20TB，与生产存储逻辑隔离。", "设备核验"],
    ["K-11", "★", "备份恢复", "至少每日增量、每周全量；定期完成数据库、对象文件和配置恢复演练，并记录 RPO/RTO。", "恢复演练"],
    ["K-12", "▲", "高可用", "应用/API 至少双节点，任务队列和模型网关避免单点；数据库、存储可按甲方条件配置主备或快照。", "故障切换"],
    ["K-13", "▲", "网络", "支持 HTTPS、固定出口 IP、代理、白名单、10GbE 校内存储网络及云端接口域名管理。", "网络核验"],
    ["K-14", "★", "兼容性", "支持主流 Chromium 内核浏览器；服务器端支持主流 Linux 或国产兼容环境，具体适配清单合同确认。", "兼容测试"],
    ["K-15", "▲", "监控告警", "监控 CPU、内存、磁盘、数据库、存储、队列、接口、模型、Tokens、费用、失败率和备份；支持阈值告警。", "监控演示"],
    ["K-16", "▲", "运营服务", "提供 7×8 服务；P1 首次响应建议不高于 30 分钟，P2 不高于 2 小时，具体 SLA 合同确认。", "服务记录"],
    ["K-17", "▲", "升级策略", "支持灰度、回滚、数据库迁移、模型版本切换和维护通知；升级不得破坏历史课程与证据。", "升级演练"],
    ["K-18", "▲", "成本治理", "支持按模型、组织、课程、功能、账号、时间统计输入/输出 Tokens、多媒体次数、缓存、失败和费用，提供四级预算告警。", "报表核验"],
]
param_table(doc, technical_rows)

# 13
page_break(doc)
heading(doc, "十三、实施交付与验收参数")
delivery_rows = [
    ["D-01", "★", "需求与原型", "完成业务调研、现有资产盘点、需求规格、角色权限、课程配置和高保真原型确认。", "签字文件"],
    ["D-02", "★", "生产系统", "交付可运行的前端、服务端、数据库、对象存储、向量索引、模型网关、任务队列和管理端。", "部署验收"],
    ["D-03", "★", "课程初始化", "将 48 课时大纲配置为 16 模块、24 次课，并形成资源分类、任务模板、量规、评测和成果包结构。", "课程验收"],
    ["D-04", "▲", "数据与接口", "交付数据字典、接口清单、模型适配配置、统一身份预留接口、导入导出和备份恢复配置。", "文档核验"],
    ["D-05", "▲", "课程资源导入", "指导甲方导入校本大纲、讲义、案例、习题、量规和制度；提供模板、校验和异常处理。", "导入验收"],
    ["D-06", "★", "功能测试", "覆盖教师教学、成长、学生学习、管理运行、课程隔离、权限、证据、补学、预警和审计主流程。", "测试报告"],
    ["D-07", "★", "并发测试", "模拟 80 名教师同时使用 AI；平台无崩溃、无数据串扰、请求不丢失，排队和降级可见。", "压力报告"],
    ["D-08", "★", "安全测试", "至少覆盖鉴权、越权、密钥、传输、敏感数据、提示词注入、工具权限、内容安全、生成标识和审计。", "安全报告"],
    ["D-09", "★", "备份恢复", "完成数据库、对象存储和配置的抽样恢复，并验证恢复后的权限、引用和版本一致。", "恢复报告"],
    ["D-10", "▲", "浏览器验收", "1440px 与 390px 验收关键页面；无关键遮挡、横向溢出、无响应主按钮和控制台严重错误。", "浏览器报告"],
    ["D-11", "▲", "培训交付", "至少完成管理员/运维、教师备授课、骨干教师课程建设等培训，提供课件、手册和可重复学习资料。", "培训记录"],
    ["D-12", "★", "文档交付", "至少包含安装部署、架构、接口、数据字典、管理员、教师、运维、备份恢复、安全与应急文档。", "资料清单"],
    ["D-13", "▲", "试点运行", "选择代表课程完成不少于一个完整备课—授课—测评—报告—补学或教师成长闭环试点。", "试点报告"],
    ["D-14", "▲", "问题整改", "问题按严重级别登记、分派、修复、复测和关闭；一般遗留不影响核心使用并形成双方确认计划。", "问题清单"],
    ["D-15", "★", "最终验收", "关键功能、80 人并发、数据隔离、模型治理、三年留存、备份恢复、培训交接全部通过。", "终验报告"],
]
param_table(doc, delivery_rows)
page_break(doc)
heading(doc, "建议实施阶段", level=2)
add_table(doc, ["阶段", "建议周期", "主要成果"], [
    ["需求与设计", "第 1–2 周", "调研、资产盘点、规格、原型和实施计划"],
    ["底座与硬件", "第 3–6 周", "账号权限、数据底座、环境和设备"],
    ["核心建设", "第 7–14 周", "课程、教学、测评、成长、AI 与管理能力"],
    ["集成测试", "第 15–17 周", "资源导入、接口、安全、并发和恢复"],
    ["试点验收", "第 18–20 周", "试点、培训、整改、交接和终验"],
], [3.5, 3.0, 10.7], font_size=8.2)

# 14
heading(doc, "十四、关键参数汇总与采购建议")
summary_rows = [
    ["1", "平台定位", "教、学、测、评、练、研、管一体化 Learning OS"],
    ["2", "核心用户", "教师、学生、教学管理者；教师具有教学工作/我的成长双空间"],
    ["3", "课程基线", "48 课时、16 模块、24 次课、8 类结课成果"],
    ["4", "教学设计", "任务驱动、项目贯穿、证据中心、掌握学习、真实性评价、人机协同"],
    ["5", "前沿内容", "上下文工程、Agent Skill、RAG、MCP、状态记忆、智能体、Harness、Agent Eval"],
    ["6", "教师能力", "5 个计分维度 + 3 项安全责任通关门槛，L1–L4 等级"],
    ["7", "实训体系", "13 类共享 AI 工具、11 项平台实训、18 项教师真实工作任务"],
    ["8", "可信治理", "来源、版本、权限、匿名、置信度、证据充分度、人工复核、生成标识和审计"],
    ["9", "建设规模", "100 名教师账号、80 人峰值 AI 并发、应用侧≥100 路在途请求"],
    ["10", "资源容量", "建议 150–200 万 TPM、三年约 65 亿等价 Tokens、10 万图/8,000 视频控制目标"],
    ["11", "数据存储", "内容留存≥3 年、对象存储建议≥10TB、独立备份建议≥20TB"],
    ["12", "部署方式", "校内业务与数据 + 云端模型弹性调用，模型与工具可替换"],
]
add_table(doc, ["序号", "维度", "关键参数"], summary_rows, [1.2, 4.0, 12.0], font_size=8.1)
callout(doc, "采购建议：核心参数应优先验证“真实闭环、权限、证据、回流和自动化验收”，避免只按页面数量或 AI 按钮数量评价。供应商演示应使用同一课程、同一教师、同一任务和同一证据贯穿关键流程。")
heading(doc, "建议供应商现场演示流程", level=2)
bullets(doc, [
    "教师基于标准课程生成并修改情境化方案，完成资源入库、检索测试及有引用、懂边界的 AI 助教发布。",
    "完成智能出题、组卷、AI 批改、教师复核、报告补学回流，以及教师摸底—培训—实训—研究—复测成长闭环。",
    "管理者完成匿名预警交办与结果回流，并模拟 80 人 AI 并发、上游限流、模型切换、低置信度和备份恢复。",
])

# 15
page_break(doc)
heading(doc, "十五、参考框架、标准与技术口径")
text(doc, "本参数以国内教育和生成式 AI 合规要求为底线，结合国际教师 AI 能力、风险治理和智能体工程公开框架。引用仅用于确定能力方向和技术口径，不构成对单一厂商或实现路线的绑定。", size=9.2, color=MUTED)
references = [
    ("JY/T 0646—2022《教师数字素养》", "https://www.moe.gov.cn/srcsite/A16/s3342/202302/W020230214594527529113.pdf", "教师数字意识、知识技能、应用、责任与专业发展基础。"),
    ("UNESCO AI Competency Framework for Teachers (2024)", "https://www.unesco.org/en/articles/ai-competency-framework-teachers", "人本、AI 伦理、基础与应用、AI 教学法、专业学习，以及 Acquire–Deepen–Create 进阶。"),
    ("NIST AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework", "以治理、映射、测量和管理组织可信 AI 风险。"),
    ("NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf", "生成式 AI 风险、测试、来源、事件和全生命周期治理。"),
    ("《生成式人工智能服务管理暂行办法》", "https://www.cac.gov.cn/2023-07/13/c_1690898327029107.htm", "生成式 AI 服务使用、个人信息、内容安全、用户权利和稳定服务要求。"),
    ("《人工智能生成合成内容标识办法》", "https://www.nrta.gov.cn/art/2025/3/14/art_113_70340.html", "生成合成内容显式/隐式标识和发布声明要求。"),
    ("GB 45438—2025《网络安全技术 人工智能生成合成内容标识方法》", "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=F32EA2A561F1886CD8D606513512D547", "生成内容标识技术方法。"),
    ("Agent Skills Specification", "https://agentskills.io/specification", "SKILL.md、scripts、references、assets 和渐进式披露结构。"),
    ("Model Context Protocol Specification 2025-06-18", "https://modelcontextprotocol.io/specification/2025-06-18/basic/index", "JSON-RPC、生命周期、授权、资源、提示和工具等协议口径。"),
    ("Google Developer’s Guide to AI Agent Protocols (2026)", "https://developers.googleblog.com/en/developers-guide-to-ai-agent-protocols/", "MCP、A2A 等智能体协议的职责边界和互操作方向。"),
    ("OpenAI Harness Engineering (2026)", "https://openai.com/index/harness-engineering/", "环境、知识、工具、可观测性、验证和反馈闭环等 Harness 工程思想。"),
]
ref_table = doc.add_table(rows=1, cols=3)
ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref_table.autofit = False
set_table_borders(ref_table)
for i, (h, w) in enumerate(zip(["参考文件", "链接", "本参数采用的重点"], [5.0, 5.0, 7.2])):
    c = ref_table.rows[0].cells[i]
    set_cell_width(c, w)
    set_cell_shading(c, TABLE_HEADER)
    set_cell_margins(c, 90, 75, 90, 75)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_run(r, 7.7, True, INK)
set_repeat_header(ref_table.rows[0])
for idx, (name, url, note) in enumerate(references):
    row = ref_table.add_row()
    prevent_split(row)
    for c in row.cells:
        set_cell_margins(c)
        set_cell_shading(c, WHITE)
    set_cell_width(row.cells[0], 5.0)
    set_cell_width(row.cells[1], 5.0)
    set_cell_width(row.cells[2], 7.2)
    p = row.cells[0].paragraphs[0]
    r = p.add_run(name)
    set_run(r, 7.2, False, INK)
    p = row.cells[1].paragraphs[0]
    add_hyperlink(p, "查看官方文件", url)
    p = row.cells[2].paragraphs[0]
    r = p.add_run(note)
    set_run(r, 7.2, False, INK)
callout(doc, "版本维护：前沿技术和监管要求变化较快，课程组与平台运营方应至少每学年复核一次模型、协议、案例、标识、安全规则和评测集；不得把单一厂商或某一版本框架的操作步骤固化为长期课程目标。", fill=AMBER_SOFT, border=AMBER)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
