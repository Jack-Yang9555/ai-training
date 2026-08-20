from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "实训资料" / "产品与交付" / "AI通识培训与教学应用平台技术参数说明书.docx"

BLUE = "1F4E79"
HEADER_FILL = "D9D9D9"
TOTAL_FILL = "EDEDED"
LIGHT_FILL = "F7F7F7"
WHITE = "FFFFFF"
LINE = "BFBFBF"
LINE_DARK = "A6A6A6"
INK = RGBColor(0, 0, 0)
BLUE_RGB = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)


def set_run(run, size=9, bold=False, color=INK, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_paragraph(paragraph, before=0, after=2, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, value="", size=9, bold=False, color=INK, align=None, before=0, after=3, line=1.35):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(value)
    set_run(run, size, bold, color)
    set_paragraph(paragraph, before, after, line)
    return paragraph


def set_paragraph_border(paragraph, edge="bottom", color=LINE_DARK, size=4, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:color"), color)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))


def add_heading(doc, value, level=1):
    paragraph = doc.add_paragraph()
    size = 18 if level == 1 else 11.5
    color = BLUE_RGB if level == 1 else INK
    run = paragraph.add_run(value)
    set_run(run, size, True, color)
    set_paragraph(paragraph, 4 if level == 1 else 6, 8 if level == 1 else 4, 1.15)
    if level == 1:
        set_paragraph_border(paragraph)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullets(doc, items, size=8.8):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run(run, size, False, INK)
        set_paragraph(paragraph, 0, 1, 1.25)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=75, start=70, bottom=75, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color=LINE, size=4, outer_color=LINE_DARK, outer_size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:color"), outer_color if edge in ("top", "left", "bottom", "right") else color)
        node.set(qn("w:sz"), str(outer_size if edge in ("top", "left", "bottom", "right") else size))


def add_table(doc, headers, rows, widths, font_size=7.5, total_rows=None, center_cols=None):
    total_rows = set(total_rows or [])
    center_cols = set(center_cols or [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    prevent_row_split(header)
    for index, (label, width) in enumerate(zip(headers, widths)):
        cell = header.cells[index]
        set_cell_width(cell, width)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell, 90, 65, 90, 65)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(label))
        set_run(run, font_size, True, INK)
        set_paragraph(paragraph, 0, 0, 1.05)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        is_total = row_index in total_rows
        for index, (value, width) in enumerate(zip(values, widths)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            set_cell_shading(cell, TOTAL_FILL if is_total else WHITE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            if index in center_cols:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            color = BLUE_RGB if index == 3 and str(value) == "★" else INK
            set_run(run, font_size, is_total or (index == 3 and str(value) == "★"), color)
            set_paragraph(paragraph, 0, 0, 1.15)
    add_text(doc, "", after=1)
    return table


def add_parameter_table(doc, rows, font_size=7.2):
    return add_table(
        doc,
        ["编号", "技术参数项", "技术指标与要求", "属性", "验收方式"],
        rows,
        [1.25, 3.1, 9.85, 0.8, 2.2],
        font_size=font_size,
        center_cols={0, 3},
    )


def add_note(doc, value):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=LINE, size=4, outer_color=LINE, outer_size=4)
    cell = table.cell(0, 0)
    set_cell_width(cell, 17.2)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, 110, 130, 110, 130)
    run = cell.paragraphs[0].add_run(value)
    set_run(run, 8.7, False, INK)
    set_paragraph(cell.paragraphs[0], 0, 0, 1.3)
    add_text(doc, "", after=1)


def page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    set_paragraph(paragraph, 0, 0, 1)


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def set_page_number_start(section, start=1):
    section_pr = section._sectPr
    page_number = section_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_pr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def add_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AI 通识培训与教学应用平台技术参数说明书")
    set_run(run, 7.8, False, MUTED)
    set_paragraph(header, 0, 3, 1.0)
    set_paragraph_border(header)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—  第 ")
    set_run(run, 8, False, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run(" 页  —")
    set_run(run, 8, False, MUTED)


def add_hyperlink(paragraph, label, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "14")
    font_size_cs = OxmlElement("w:szCs")
    font_size_cs.set(qn("w:val"), "14")
    run_props.extend([fonts, font_size, font_size_cs, color, underline])
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = label
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
cover_section = doc.sections[0]
configure_page(cover_section)
cover_section.top_margin = Cm(1.9)
cover_section.bottom_margin = Cm(1.9)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(9)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.3
for style_name in ("List Bullet", "List Number"):
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(8.8)

doc.core_properties.title = "AI通识培训与教学应用平台技术参数说明书"
doc.core_properties.subject = "高职院校采购技术需求、供应商响应与验收基线"
doc.core_properties.author = "启境 AI Learning OS 项目组"

# 封面
add_text(doc, "TECHNICAL PARAMETER SPECIFICATION", size=10.5, bold=True, color=BLUE_RGB, align=WD_ALIGN_PARAGRAPH.CENTER, before=42, after=22)
add_text(doc, "AI 通识培训与教学应用平台", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=5, line=1.1)
add_text(doc, "技术参数说明书", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14, line=1.1)
add_text(doc, "适用于项目立项、采购技术需求、供应商响应及履约验收", size=11.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)

cover_rows = [
    ["文档版本", "V2.0", "编制日期", "2026 年 7 月 31 日"],
    ["采购规模", "100 名教师正式账号", "AI 峰值并发", "不少于 80 人"],
    ["课程基线", "48 课时 / 16 模块 / 24 次课", "部署模式", "校内部署 + 云端模型"],
    ["数据留存", "不少于 3 年", "服务期限", "三年 7×8 服务"],
]
add_table(doc, ["项目", "技术口径", "项目", "技术口径"], cover_rows, [2.3, 6.3, 2.3, 6.3], font_size=8.4)
add_note(doc, "本说明书规定项目最低技术指标、功能边界、容量要求、交付成果与验收方式。供应商应逐项响应，不得以产品介绍、演示截图或笼统承诺替代技术响应。")

body_section = doc.add_section(WD_SECTION.NEW_PAGE)
configure_page(body_section)
add_header_footer(body_section)
set_page_number_start(body_section, 1)

# 一、范围与口径
add_heading(doc, "一、项目范围与参数口径")
add_heading(doc, "1. 建设范围", level=2)
add_bullets(doc, [
    "本期建设对象为 100 名教师，共 100 个正式账号；平台管理权限由上述账号中的授权人员承担，满足不少于 80 名教师同时发起 AI 任务。",
    "建设内容包括软件平台、混合部署、模型资源、必要硬件扩容、课程初始化、实施培训、三年运维和验收交付。",
    "学生账号及学生直接调用 AI 不纳入本期账号容量和模型资源验收；平台架构应预留后续学生角色扩展能力。",
    "本期以 48 课时 AI 通识课程为课程基线，形成可配置、可教学、可实训、可测评和可追溯的课程运行单元。",
])
add_heading(doc, "2. 参数使用规则", level=2)
legend_rows = [
    ["★", "差异化核心参数", "标识平台特色、关键闭环或核心治理能力，供应商应提供功能演示、界面截图、检测报告或等效证明。"],
    ["空白", "一般技术参数", "属于本项目最低技术要求，供应商仍须逐项响应并纳入合同验收。"],
    ["技术指标与要求", "强制响应口径", "“应、必须、不低于、至少”等表述均为最低要求；出现数量、容量、时限或状态时应按该指标验收。"],
]
add_table(doc, ["标识", "性质", "响应规则"], legend_rows, [2.0, 4.0, 11.2], font_size=8.2)
add_note(doc, "总体建设目标、教学理念和课程内容不作为孤立参数编号；只有能够形成系统功能、配置项、容量指标、交付物或验收证据的要求进入技术参数表。")

# 二、软件平台
add_heading(doc, "二、软件平台技术参数")

add_heading(doc, "1. 总体架构与平台底座", level=2)
architecture_rows = [
    ["A-01", "混合部署架构", "业务系统、账号权限、关系数据库、对象存储、向量索引、审计日志和模型网关应部署于校内或甲方指定环境；经批准的文本、视觉、图片、音视频模型可通过云端 API 调用。", "★", "架构图、部署核验"],
    ["A-02", "分层架构", "系统至少应包含用户访问层、教学业务层、AI 能力层、数据与知识层、集成层、安全治理层和基础设施层，各层接口和责任边界清晰。", "", "设计文档"],
    ["A-03", "统一模型网关", "所有模型调用必须经服务端统一网关；网关应支持模型路由、限流、超时、重试、熔断、配额、成本统计和供应商切换，浏览器端不得保存模型密钥。", "★", "接口与故障测试"],
    ["A-04", "异步任务中心", "图片、视频、批量报告、知识库解析和批量批改等长任务应异步执行，并提供进度、排队、取消、失败原因、重试和完成通知。", "", "任务演示"],
    ["A-05", "统一知识底座", "课程大纲、讲义、案例、习题、量规、任务单和制度规范应形成统一知识源，至少记录来源、版本、适用模块、知识点、权限和更新时间。", "★", "数据抽查"],
    ["A-06", "AI 服务适配", "文本模型、视觉理解、图片生成、视频生成、Embedding、Rerank、OCR、文档解析和内容安全能力应通过适配层接入并可替换。", "", "替换演示或接口文档"],
    ["A-07", "工作流与智能体编排", "应支持配置化编排模型、知识、Skill、工具、状态、规则、护栏和人工确认节点，运行轨迹应可查看、检索和导出。", "★", "场景演示"],
    ["A-08", "开放接口", "应提供标准 REST API；工具连接应预留 MCP 等开放协议，多智能体协作和 A2A 作为扩展能力。", "", "接口清单"],
    ["A-09", "数据主权", "甲方课程资料、提示词、生成内容、作品、评分、报告、日志和研究成果归甲方所有，系统应支持导出、迁移、归档和按权限删除。", "★", "合同与数据测试"],
    ["A-10", "可观测性", "应监控业务接口、模型调用、任务队列、耗时、Tokens、费用、失败率、低置信度、人工接管、存储和备份状态。", "", "监控演示"],
]
add_parameter_table(doc, architecture_rows)

add_heading(doc, "2. 账号、权限与工作空间", level=2)
account_rows = [
    ["U-01", "教师账号容量", "首期应支持不少于 100 名教师正式账号，并支持按院系、专业、课程和组织层级管理。", "", "账号测试"],
    ["U-02", "管理权限", "应支持在 100 个正式账号中授权平台管理员、教学管理者和运维管理员等管理角色，并允许按职责配置菜单、对象、数据和操作权限。", "", "角色测试"],
    ["U-03", "教师双空间", "教师端必须提供“教学工作”和“我的成长”两个独立工作空间；教学班级上下文与个人课程、笔记、实训和能力报告不得相互覆盖。", "★", "状态与数据测试"],
    ["U-04", "教师工作台", "教师教学工作空间至少应包含智能教学、学习促进、测评管理、教学诊断、AI 研究和自由实训六类任务入口。", "★", "角色演示"],
    ["U-05", "成长工作台", "教师成长空间至少应包含摸底测评、能力报告、培训计划、课程学习、自主路径、AI 实训、研究成果和复测入口。", "★", "全流程演示"],
    ["U-06", "集中权限策略", "入口权限、对象权限、数据权限和状态调用权限必须由服务端集中策略控制，前端隐藏不得代替服务端鉴权。", "", "接口越权测试"],
    ["U-07", "课程与对象隔离", "课程、班级、任务、测评、报告和补学状态应按 courseId 或等效主键隔离；切换对象不得出现数据串扰。", "", "数据隔离测试"],
    ["U-08", "组织范围", "应支持学校、院系、专业、课程和教师组等组织范围配置，并支持授权范围继承与回收。", "", "配置演示"],
    ["U-09", "统一身份预留", "应预留标准统一身份认证接口，至少支持 OAuth2、OIDC、CAS 或甲方现有统一认证方式之一，具体适配在合同附件确认。", "", "接口文档"],
    ["U-10", "无障碍与兼容操作", "关键页面应支持键盘操作、清晰焦点、语义标题、替代文本和合理字号；移动端不得出现关键遮挡或横向溢出。", "", "可访问性测试"],
]
add_parameter_table(doc, account_rows)

add_heading(doc, "3. 课程建设与智能教学", level=2)
teaching_rows = [
    ["T-01", "课程结构", "应将 AI 通识课程配置为 16 个模块、24 次课、48 课时；每次课至少包含目标、知识点、流程、资源、活动、任务、题目、量规和课后安排。", "★", "课程结构验收"],
    ["T-02", "智能备课流程", "教师备课流程至少应包含标准方案选择、情境配置、AI 方案生成、方案对比、人工修改、采纳发布和版本留痕。", "★", "全流程演示"],
    ["T-03", "教学情境配置", "应支持教学对象、基础水平、人数、授课方式、课时、难度和教学重点等情境参数配置。", "", "配置演示"],
    ["T-04", "适配依据", "AI 生成方案必须展示标准方案与适配方案在目标、难度、节奏、活动、资源和题目方面的差异及依据。", "★", "结果核验"],
    ["T-05", "人工采纳与发布", "AI 方案应先保存为草稿；教师修改并确认后方可发布。系统必须同时保留 AI 原稿、人工修改、发布版本和操作人。", "★", "版本测试"],
    ["T-06", "配套成果生成", "采纳方案后应至少形成情境化教案、课堂活动包、配套题目集和教学进度方案四类结构化成果。", "", "成果抽查"],
    ["T-07", "教学实施", "课堂实施应支持步骤时长调整、资源版本替换、题目增删、知识点查看、准备度检查和任务发布。", "", "课堂演示"],
    ["T-08", "课程资源管理", "资源应至少分为标准资源、AI 生成资源和校本资源，支持来源、版本、状态、权限、标签、知识点和引用关系管理。", "", "资源核验"],
    ["T-09", "资源导入", "应支持大纲、讲义、案例、习题、量规、任务单和术语表导入，并记录解析、分段、重复率、覆盖率和异常信息。", "", "导入演示"],
    ["T-10", "知识图谱", "应支持从课程知识库生成知识图谱，展示知识点、先修关系、资源、题目、任务和测评之间的关联。", "", "图谱演示"],
    ["T-11", "AI 助教配置", "AI 助教配置至少应包含知识绑定、教学人格、呈现形态、测试和发布五个环节；知识库为必选底座，知识图谱为可选增强。", "★", "配置演示"],
    ["T-12", "有引用回答", "AI 助教回答应展示课程来源、章节、知识点和置信度；超出知识范围时应明确知识边界并提示教师确认。", "★", "问答测试"],
    ["T-13", "助教发布", "助教发布前应完成测试集、版本、开放范围、课程入口、引用、边界、内容安全和人工接管检查。", "", "发布演示"],
    ["T-14", "助教运营", "应统计解决率、知识命中率、高频问题、低置信度回答、越界问题和人工接管情况。", "", "运营报表"],
    ["T-15", "多模态资源", "应支持图文、流程图、海报、音频、短视频或数字人脚本等资源生成；成果需记录提示词、模型、版本、人工修改、授权和 AI 标识。", "", "成果抽查"],
    ["T-16", "课程版本", "课程、模块、课次和资源应支持版本管理、发布版本与编辑草稿隔离、历史引用回溯和版本对比。", "", "版本测试"],
]
add_parameter_table(doc, teaching_rows)

add_heading(doc, "4. 学习、测评与诊断", level=2)
assessment_rows = [
    ["E-01", "学习路径", "教师成长学习应支持课程导学、继续学习、资源进度、收藏、笔记、AI 推荐路径和自主路径。", "", "学习演示"],
    ["E-02", "多周期观察", "应支持本课、本周、本单元等观察周期，并标识过程性数据与正式成绩数据。", "", "筛选演示"],
    ["E-03", "多源证据", "应按目标聚合学习任务、课堂练习、作品迭代、测评过程、AI 助教对话、实训和研究成果，显示来源、充分度和更新时间。", "★", "证据抽查"],
    ["E-04", "诊断区分", "系统必须区分“尚未掌握”和“证据不足”；低置信度、证据冲突或证据不足结果不得直接触发高影响自动决策。", "★", "规则测试"],
    ["E-05", "分层任务", "应支持补强、巩固和拓展三类任务，教师可配置目标、资源、题目、难度、时长、截止日期和接收对象。", "", "任务演示"],
    ["E-06", "补学回流", "补学任务完成后应更新掌握诊断、能力证据和下一轮学习建议，并保留前后变化。", "★", "闭环测试"],
    ["E-07", "测评场景", "应支持摸底测、随堂测、单元测、正式考试、实训和综合项目不少于六类测评场景。", "", "测评演示"],
    ["E-08", "题型", "题库至少应支持单选、多选、判断、填空、简答、案例分析和创作实操七类题型。", "", "题库演示"],
    ["E-09", "组卷蓝图", "应支持按总分、知识覆盖、难度、题型、认知层级和评分主体组卷，并检查相似题、量规和总分。", "", "组卷测试"],
    ["E-10", "多主体评分", "综合项目应支持机器客观评分、AI 作品评分和教师人工评分；默认比例 30/50/20，并允许课程级配置。", "★", "评分演示"],
    ["E-11", "评分证据", "AI 作品评分应至少展示任务符合度、专业准确性初检、方法与工具、迭代证据、内容安全和量规条目。", "★", "证据抽查"],
    ["E-12", "人工复核", "教师可调整 AI 分和人工分，最终分应实时重算；AI 原分、教师修改理由和最终分必须同时保留。", "★", "评分测试"],
    ["E-13", "报告体系", "应支持提交、成绩、达成、复核、难度、分布、知识掌握和共性误区等报告，并可下钻到原始证据。", "", "报告演示"],
    ["E-14", "过程性评价", "能力评价应综合学习过程、知识测评、核心实训作品和综合项目，不得仅以点击次数或单次测验评价。", "", "评价方案核验"],
]
add_parameter_table(doc, assessment_rows)

add_heading(doc, "5. 教师成长、实训与研究", level=2)
growth_rows = [
    ["G-01", "成长闭环", "应形成摸底测评、能力报告、培训方案、课程学习、自主学习、AI 实训、研究/工具开发、成长报告和复测的连续闭环。", "★", "全流程演示"],
    ["G-02", "5+3 能力模型", "五个计分维度至少包含 AI 基础认知、提示词与多模态、知识库与智能体、教学融合、研究创新；设置事实核验、数据版权和人工责任三类通关门槛。", "★", "报告核验"],
    ["G-03", "能力等级", "应支持 L1 入门、L2 实践、L3 创新和 L4 引领四级能力等级，等级规则可配置并有证据支撑。", "", "规则测试"],
    ["G-04", "真实摸底", "摸底测评应包含理论、情境和实操；实操题应提交提示词、截图、文件、测试记录或人工确认。", "", "作答演示"],
    ["G-05", "个性化培训计划", "应基于摸底报告从已导入课程生成培训计划，支持调整学习顺序、每周投入和选修内容。", "", "计划演示"],
    ["G-06", "成果证据台账", "课程学习、自由实训、教学实施、评分复核、分层干预、研究成果和复测等至少七类成果应进入个人证据台账并可下钻。", "★", "证据抽查"],
    ["G-07", "教师真实任务", "平台应配置不少于 18 项教师真实工作任务，覆盖备课资源、命题评价、知识库智能体和教学研究四类任务链。", "", "任务清单"],
    ["G-08", "平台实训任务", "自由实训应提供不少于 11 项平台技术实训，覆盖知识库、检索优化、生成图谱、配置智能体、发布助教、命题、组卷、批改和分析。", "", "任务演示"],
    ["G-09", "共享工具目录", "应提供不少于 13 类共享 AI 工具，覆盖研究、文献、数据、提示、多模态、视频、RAG、低代码、智能体、MCP、识别、Vibe Coding 和模型评测。", "", "工具目录"],
    ["G-10", "实训任务字段", "每项实训应至少记录频率、难度、标准时长、前置任务、输入、操作、产出、量规和完成状态。", "", "任务抽查"],
    ["G-11", "实训成果回流", "实训成果应进入教师实训记录和能力报告，记录工具、任务、用时、得分、作品、版本和人工确认。", "", "回流演示"],
    ["G-12", "AI 研究工作台", "仅教师可访问，至少包含研究项目、共享 AI 工具、科研智能体、成果管理和研究伦理。", "★", "权限测试"],
    ["G-13", "研究项目", "研究流程应支持问题设计、文献与证据、数据分析与干预、成果与复核四阶段，并保留来源、参数和人工判断。", "", "研究演示"],
    ["G-14", "科研智能体", "应支持绑定研究项目、导入资料、配置角色、工具、测试集和版本，并记录测试与成果关联。", "", "配置演示"],
    ["G-15", "延迟复测", "应支持首次摸底、过程评价、结业复测和 30/90 天延迟复测，展示能力变化和任务迁移证据。", "", "报告演示"],
]
add_parameter_table(doc, growth_rows)

# 三、前沿能力与治理
add_heading(doc, "三、前沿 AI 能力与可信治理参数")
add_heading(doc, "1. 前沿 AI 能力", level=2)
frontier_rows = [
    ["F-01", "上下文工程", "应支持目标、背景、资料、术语、示例、限制、状态、质量标准和版本组织，并检查缺失、冲突、过时、过载和污染。", "★", "实训演示"],
    ["F-02", "结构化输出", "应支持表格、JSON 或 Schema 约束，校验字段、类型、必填、异常值和失败回退；确定性计算应优先调用工具。", "", "校验测试"],
    ["F-03", "Agent Skill", "应支持按 SKILL.md、references、scripts、assets 等结构组织复用技能，并记录触发条件、输入、步骤、边界、依赖和验收标准。", "★", "成果核验"],
    ["F-04", "RAG 检索", "应支持文档切分、向量化、关键词/混合检索、重排、引用和知识边界，并提供正常、模糊、无答案和版本冲突测试。", "★", "检索测试"],
    ["F-05", "智能体定义", "智能体配置至少应包含目标、指令、上下文、知识、Skill、工具、状态、规划、权限、护栏和人工接管。", "", "配置核验"],
    ["F-06", "MCP 教学与连接", "应支持资源、提示和工具等 MCP 基本概念，在仿真或受控环境配置只读工具；不得要求所有教师开发 MCP 服务。", "", "实训演示"],
    ["F-07", "状态与记忆", "应区分当前上下文、会话记忆、业务状态和长期知识，支持状态流转、断点保存、超时重试、幂等和失败转人工。", "", "故障测试"],
    ["F-08", "Harness", "应将运行环境、上下文、知识、工具、规则、日志、反馈和恢复机制组织为可查看配置，并可关联运行轨迹。", "★", "轨迹核验"],
    ["F-09", "Agent Eval", "应支持固定评测集重复执行，记录成功率、引用、耗时、成本、工具调用、人工接管和失败类别。", "★", "评测演示"],
    ["F-10", "评测场景", "Agent Eval 至少应覆盖正常、模糊、越界、攻击、高风险和故障六类场景。", "", "评测报告"],
    ["F-11", "红队测试", "应支持提示词注入、数据泄露、越权调用、错误工具结果、失败循环和高风险请求测试，并在修订护栏后复测。", "", "安全实训"],
    ["F-12", "计算机使用智能体", "只能在隔离仿真环境执行观察、规划、填写、检查和确认；登录、验证码、敏感字段、最终提交和不可逆操作必须转人工。", "", "沙箱演示"],
    ["F-13", "多模态", "应支持文本、图片、音频和视频的生成与理解任务，成果应包含提示词、来源、版权、字幕、替代文本、人工核验和 AI 标识。", "", "成果抽查"],
    ["F-14", "模型评测", "应支持使用同一任务、输入和量规对比多模型的质量、稳定性、速度、成本、隐私和安全，并保留选择依据。", "", "评测报告"],
    ["F-15", "深度研究", "应支持研究问题、检索范围、证据矩阵、来源核验、冲突证据、过程记录和人工结论，禁止伪造引用。", "", "研究演示"],
    ["F-16", "Vibe Coding", "应支持自然语言生成轻量网页或工具原型，并保留需求、验收标准、依赖、测试、修改记录和安全检查。", "", "原型验收"],
]
add_parameter_table(doc, frontier_rows)

add_heading(doc, "2. 可信 AI、数据安全与合规", level=2)
security_rows = [
    ["S-01", "统一证据字段", "关键 AI 结果至少应记录来源、知识点、输入摘要、模型/服务版本、提示版本、量规、结论、置信度、证据充分度、权限、匿名状态、人工复核、修改记录和时间。", "★", "字段抽查"],
    ["S-02", "人工复核状态", "应支持待复核、已确认、已修改和已驳回等状态，AI 原值与最终值必须同时保留。", "★", "状态测试"],
    ["S-03", "低置信度治理", "低置信度、无引用、证据冲突或证据不足必须显式标记并进入人工处理，不得静默输出确定性结论。", "★", "规则测试"],
    ["S-04", "最小必要传输", "向云端模型只能发送完成任务所需最小数据，应支持脱敏、敏感实体检测、禁止字段和传输审计。", "", "接口抓包与日志"],
    ["S-05", "密钥安全", "模型、存储、数据库和第三方工具密钥必须由服务端管理，支持轮换、权限隔离和泄露应急，不得写入前端代码。", "", "安全检查"],
    ["S-06", "个人信息保护", "应支持个人信息查阅、更正、导出和删除流程，对教师、课程和研究数据实行分级分类与最小权限。", "", "流程测试"],
    ["S-07", "版权与授权", "资源和生成内容应记录来源、授权、适用范围、AI 参与方式、人工修改和发布确认；不明版权内容不得公开发布。", "", "内容抽查"],
    ["S-08", "生成内容标识", "导出的 AI 生成/合成文本、图片、音频和视频应按适用法规和标准提供显式提示、元数据或等效标识。", "★", "导出核验"],
    ["S-09", "内容安全", "应支持输入输出安全检测、违法违规内容处置、用户警示、功能限制、事件记录和申诉/举报流程。", "", "安全测试"],
    ["S-10", "高风险边界", "医疗、工程、财务和高影响教学决策等场景必须提示正式规范和人工确认，AI 不得直接作最终决定。", "", "情境测试"],
    ["S-11", "工具权限", "工具应按只读/写入、可逆/不可逆、低/中/高风险分级；写入、发送、支付、删除和外部发布必须人工确认。", "", "权限测试"],
    ["S-12", "审计日志", "应记录登录、权限、课程发布、模型调用、工具调用、评分修改、数据导出、删除、配置变化和异常事件，日志可检索并防篡改。", "", "审计抽查"],
    ["S-13", "匿名研究", "研究和管理数据应默认匿名或去标识，成果导出前应检查再识别风险、来源授权和 AI 使用披露。", "", "研究抽查"],
    ["S-14", "模型供应商管理", "应记录模型备案/合规信息、供应商、版本、适用场景、价格、数据策略、可用区和替代方案。", "", "台账核验"],
    ["S-15", "风险闭环", "预警应支持待处理、已交办、已解决等状态，并可从管理关注回流教师干预、任务完成和指标更新。", "★", "跨角色演示"],
    ["S-16", "AI 事件管理", "应提供模型异常、错误引用、越权、敏感信息、内容标识缺失和供应商不可用事件的登记、处置、复盘和改进。", "", "应急演练"],
]
add_parameter_table(doc, security_rows)

add_heading(doc, "3. 教学运行管理", level=2)
management_rows = [
    ["M-01", "运行总览", "应展示课程、教师、任务、测评、报告、AI 应用、预警和服务健康等聚合指标，并支持组织范围筛选。", "", "管理演示"],
    ["M-02", "课程运行", "应统计课程开设、教师、资源就绪、学习进度、测评状态和异常对象。", "", "报表核验"],
    ["M-03", "学习质量", "应统计目标达成、证据充分度、补学完成、作品迭代和知识点趋势，不得展示无授权实名。", "", "报表核验"],
    ["M-04", "测评运行", "应统计测评状态、提交、批改、人工复核、争议和异常任务。", "", "报表核验"],
    ["M-05", "AI 应用治理", "应展示 Tokens、费用、模型、成功率、响应、引用命中、低置信度、人工接管、内容安全和工具调用指标。", "★", "运营演示"],
    ["M-06", "教师成长管理", "应仅展示授权范围内的培训参与、任务完成、能力分布、成果类型和支持需求，不得查看个人答案、笔记和研究私密内容。", "", "权限测试"],
    ["M-07", "成果资产", "应支持课程模板、提示词、量规、知识库、Skill、智能体、评测集和研究成果的归档、审核、复用和版本管理。", "", "资产演示"],
    ["M-08", "预警干预", "预警应包含对象范围、证据、责任人、截止时间、状态和干预结果，并支持交办、跟踪、解决和指标回流。", "★", "闭环演示"],
    ["M-09", "质量指标", "至少应支持摸底/结业变化、任务达成、人工修改质量、引用命中、工具成功、人工接管、安全通关和满意度指标。", "", "指标核验"],
    ["M-10", "数据导出", "应支持按学期或年度导出课程运行、教师发展、AI 成本、风险事件、成果资产和改进计划。", "", "导出测试"],
]
add_parameter_table(doc, management_rows)

# 四、容量与硬件
add_heading(doc, "四、性能、容量、部署与硬件参数")
capacity_rows = [
    ["C-01", "教师账号", "首期正式教师账号容量不得低于 100 个。", "", "账号测试"],
    ["C-02", "AI 峰值并发", "平台应支持不少于 80 名教师同时发起 AI 任务；应用网关在途请求能力不得低于 100 路。", "★", "压力测试"],
    ["C-03", "文本服务配额", "企业文本模型可用配额不得低于 150 万 TPM，目标配置为 200 万 TPM；应通过队列和路由控制智能体扇出。", "", "配额核验与压测"],
    ["C-04", "模型资源容量", "三年文本模型资源不得低于 65 亿 Tokens（按等价文本 Token 计），应包含系统提示、重试、工具链和增长余量。", "★", "资源台账"],
    ["C-05", "多媒体资源", "三年资源池应满足不少于 10 万次图片生成、8,000 次短视频生成，或提供经甲方确认的等值预算池。", "", "配额核验"],
    ["C-06", "业务响应", "校内网络正常条件下，普通业务页面 95% 请求响应时间不得高于 3 秒。", "", "性能报告"],
    ["C-07", "AI 任务响应", "文本任务应支持流式返回；长任务进入可见队列，用户可查看状态、取消、超时、重试和失败原因。", "", "接口测试"],
    ["C-08", "内容留存", "课程资料、生成内容、作品、评分、报告和关键审计数据保存期限不得少于 3 年。", "★", "数据抽查"],
    ["C-09", "对象存储", "可用对象存储容量不得低于 10TB，并支持版本、校验、生命周期、归档和按权限下载。", "", "容量核验"],
    ["C-10", "备份容量", "独立备份可用容量不得低于 20TB，并与生产存储逻辑隔离。", "", "设备核验"],
    ["C-11", "备份策略", "至少应执行每日增量、每周全量备份，并完成数据库、对象文件和配置恢复演练。", "", "恢复演练"],
    ["C-12", "应用高可用", "应用/API 应采用不少于 2 个节点，任务队列和模型网关不得存在单点；数据库和存储按甲方条件配置主备或快照。", "", "故障切换"],
    ["C-13", "网络", "应支持 HTTPS、固定出口 IP、代理、白名单、10GbE 校内存储网络和云端接口域名管理。", "", "网络核验"],
    ["C-14", "服务器兼容", "服务器端应支持主流 Linux 或国产兼容环境，具体操作系统、数据库和中间件适配清单在合同附件确认。", "", "兼容测试"],
    ["C-15", "浏览器兼容", "应支持主流 Chromium 内核浏览器；1440px 和 390px 关键页面不得出现遮挡、横向溢出或无响应主操作。", "", "浏览器验收"],
    ["C-16", "监控告警", "应监控 CPU、内存、磁盘、数据库、存储、队列、接口、模型、Tokens、费用、失败率和备份状态，并支持阈值告警。", "", "监控演示"],
    ["C-17", "成本治理", "应按模型、组织、课程、功能、账号和时间统计输入/输出 Tokens、多媒体次数、缓存、失败和费用，并提供个人、课程、月度和项目总额四级预算告警。", "", "报表核验"],
]
add_parameter_table(doc, capacity_rows)

add_heading(doc, "硬件最低配置", level=2)
hardware_rows = [
    ["H-01", "应用/API 计算节点", "机架式；物理核心不少于 24；内存不少于 128GB ECC；2×1.92TB 企业级 SSD RAID1；双 10GbE；冗余电源；三年质保。", "2 台", "设备到货与性能核验"],
    ["H-02", "对象存储节点", "不少于 8 盘位；原始容量不少于 48TB；支持 RAID6 或纠删码；可用容量不少于 10TB；SSD 缓存；10GbE；快照和校验。", "1 套", "容量与功能核验"],
    ["H-03", "独立备份节点", "原始容量不少于 32TB；可用容量不少于 20TB；支持增量、全量、保留策略和恢复校验；与生产存储逻辑隔离。", "1 套", "容量与恢复核验"],
    ["H-04", "网络及安装配件", "包含 10GbE 光模块/网卡、线缆、磁盘备件、导轨、电源和安装调试材料。", "1 批", "清单与安装核验"],
]
add_table(doc, ["编号", "设备", "最低技术配置", "数量", "验收方式"], hardware_rows, [1.2, 3.0, 9.7, 1.2, 2.1], font_size=7.3, center_cols={0, 3})
add_note(doc, "硬件采购前应完成甲方现有资产盘点。现有设备达到最低配置时，可按双方书面变更用于存储、备份、安全或模型资源扩容，但不得降低本说明书规定的容量、性能和验收标准。")

# 五、实施交付
add_heading(doc, "五、实施、交付、培训与运维参数")
delivery_rows = [
    ["D-01", "需求与设计", "应完成业务调研、现有资产盘点、需求规格、角色权限、课程配置、高保真原型和实施计划确认。", "", "签字文件"],
    ["D-02", "生产系统", "应交付可运行的前端、服务端、数据库、对象存储、向量索引、模型网关、任务队列和管理端。", "★", "部署验收"],
    ["D-03", "课程初始化", "应将 48 课时大纲配置为 16 模块、24 次课，并形成资源分类、任务模板、量规、评测框架和成果包结构。", "★", "课程验收"],
    ["D-04", "课程资源边界", "本期应交付课程结构、模板、量规、评测框架和示例资源；甲方完整校本讲义、完整课程视频及学术版权终审不属于平台技术交付。", "", "交付清单核验"],
    ["D-05", "数据与接口", "应交付数据字典、接口清单、模型适配配置、统一身份预留接口、导入导出和备份恢复配置。", "", "文档核验"],
    ["D-06", "资源导入", "应指导甲方导入校本大纲、讲义、案例、习题、量规和制度，并提供模板、校验和异常处理。", "", "导入验收"],
    ["D-07", "功能测试", "应覆盖教师教学、成长、学习、测评、研究、管理运行、课程隔离、权限、证据、补学、预警和审计主流程。", "", "测试报告"],
    ["D-08", "并发测试", "应模拟 80 名教师同时使用 AI；平台不得崩溃、不得出现数据串扰或请求丢失，排队和降级状态必须可见。", "★", "压力报告"],
    ["D-09", "安全测试", "至少应覆盖鉴权、越权、密钥、传输、敏感数据、提示词注入、工具权限、内容安全、生成标识和审计。", "", "安全报告"],
    ["D-10", "备份恢复", "应完成数据库、对象存储和配置的抽样恢复，并验证恢复后的权限、引用和版本一致。", "", "恢复报告"],
    ["D-11", "浏览器验收", "应在 1440px 和 390px 验收关键页面，无关键遮挡、横向溢出、无响应主按钮和控制台严重错误。", "", "浏览器报告"],
    ["D-12", "培训交付", "平台管理员与运维培训不少于 2 场、每场不少于 3 小时；教师应用培训不少于 4 场、每场不少于 3 小时；骨干教师工作坊不少于 2 场、每场不少于 3 小时。", "", "培训记录"],
    ["D-13", "文档交付", "至少应包含安装部署、架构、接口、数据字典、管理员、教师、运维、备份恢复、安全和应急文档。", "", "资料清单"],
    ["D-14", "试点运行", "应选择代表课程完成不少于 1 个完整备课、授课、测评、报告和补学闭环试点。", "", "试点报告"],
    ["D-15", "问题整改", "问题应按严重级别登记、分派、修复、复测和关闭；一般遗留不得影响核心使用，并形成双方确认计划。", "", "问题清单"],
    ["D-16", "建设周期", "项目建设周期为合同生效且甲方提供必要环境、账号和资料后的 20 周。", "", "实施计划"],
    ["D-17", "最终验收", "关键功能、80 人并发、数据隔离、模型治理、三年留存、备份恢复、培训和资料交接全部通过后方可终验。", "★", "终验报告"],
]
add_parameter_table(doc, delivery_rows)

add_heading(doc, "三年运维服务", level=2)
service_rows = [
    ["O-01", "服务方式", "提供三年 7×8 技术服务，具体服务时段和法定节假日安排在合同中确认。", "服务记录"],
    ["O-02", "P1 紧急事件", "首次响应不高于 30 分钟，4 小时内提供恢复或绕行方案，并持续跟踪至恢复。", "工单记录"],
    ["O-03", "P2 高级事件", "首次响应不高于 2 小时，1 个服务日内提供修复或可用替代方案。", "工单记录"],
    ["O-04", "P3 一般事件", "首次响应不高于 4 小时，3 个服务日内处理或进入最近维护版本。", "工单记录"],
    ["O-05", "P4 咨询事件", "首次响应不高于 1 个服务日，按双方确认计划答复或安排。", "工单记录"],
    ["O-06", "服务内容", "应包含故障响应、巡检、备份核验、小版本升级、模型适配、使用报告和年度健康检查。", "服务报告"],
    ["O-07", "升级与回滚", "应支持灰度发布、版本回滚、数据库迁移、模型版本切换和维护通知，升级不得破坏历史课程与证据。", "升级演练"],
]
add_table(doc, ["编号", "服务参数项", "服务指标与要求", "验收方式"], service_rows, [1.2, 3.6, 10.2, 2.2], font_size=7.7, center_cols={0})

# 六、课程附件
add_heading(doc, "六、48 课时课程内容与平台能力映射")
add_text(doc, "本章用于验证平台对课程内容的承载能力，不替代软件技术参数。课程模块应能够配置为课程目标、知识点、资源、任务、工具、成果、量规和风险边界。", size=9, color=MUTED, after=6)
course_rows = [
    ["CM-01", "AI 基础、职业变革与人机分工", "4", "课程导学、课前诊断、学习路径、深度研究/模型评测", "AI 应用辨识图、岗位任务人机分工表"],
    ["CM-02", "数据、机器学习与模型评价", "4", "随堂学、研究数据分析室、模型评测场、单元测", "数据质量诊断表、模型选择说明"],
    ["CM-03", "生成式 AI 原理与可信使用", "4", "按需资源、模型评测、风险情境测评、成长报告", "AI 输出核验记录、可信使用清单"],
    ["CM-04", "上下文工程", "2", "提示词实验室、AI 助教、作品版本", "岗位任务上下文 v1/v2"],
    ["CM-05", "提示词与结构化输出", "4", "提示词实验室、数据分析、模型评测", "三轮提示词记录、结构化输出模板"],
    ["CM-06", "Agent Skill 设计", "2", "提示词实验室、Vibe Coding、成果模板", "入门级 SKILL.md、试用反馈"],
    ["CM-07", "多模态内容生成", "2", "多模态工坊、识别、作品版本", "多模态作品、版权和无障碍检查"],
    ["CM-08", "语音与数字人", "2", "文生视频、多模态工坊、数字人形态", "语音或数字人方案、字幕与披露记录"],
    ["CM-09", "计算机使用智能体", "2", "仿真页面、智能体搭建、工具连接", "受控操作轨迹、人工确认清单"],
    ["CM-10", "课程知识库与 RAG", "4", "课程资源、知识库 RAG、AI 助教、单元实训", "来源清单、6—9 题检索测试报告"],
    ["CM-11", "工具调用与 MCP", "2", "MCP 工具连接台、低代码应用、权限矩阵", "工具说明书、输入输出样例、权限表"],
    ["CM-12", "状态、记忆与任务恢复", "2", "智能体搭建、低代码应用、运行轨迹", "状态流转图、断点恢复和幂等测试"],
    ["CM-13", "智能体设计与低代码开发", "4", "智能体搭建、低代码应用、MCP、护栏", "智能体原型、权限风险矩阵、人机交接"],
    ["CM-14", "Harness 与运行可观测性", "2", "Vibe Coding、模型评测、轨迹与日志", "Harness 画布、运行轨迹标注"],
    ["CM-15", "智能体评测与安全", "4", "模型评测、MCP、红队测试、评分证据", "不少于 8 例评测集、红队报告、原型 v2"],
    ["CM-16", "专业智能体综合项目", "4", "作品版本、模型评测、成长报告、答辩", "AI 职业应用成果包、验收和治理清单"],
    ["", "合计", "48", "16 个模块 / 24 次课", "完整 AI 职业应用成果包"],
]
add_table(doc, ["编号", "模块主题", "课时", "平台能力映射", "核心成果"], course_rows, [1.2, 4.2, 1.0, 6.0, 4.8], font_size=7.25, total_rows={16}, center_cols={0, 2})

# 七、规范性依据
add_heading(doc, "七、规范性依据与版本管理")
references = [
    ("JY/T 0646—2022《教师数字素养》", "https://www.moe.gov.cn/srcsite/A16/s3342/202302/W020230214594527529113.pdf", "教师数字意识、知识技能、应用、责任与专业发展。"),
    ("UNESCO AI Competency Framework for Teachers (2024)", "https://www.unesco.org/en/articles/ai-competency-framework-teachers", "人本、AI 伦理、基础应用、AI 教学法、专业学习和能力进阶。"),
    ("NIST AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework", "可信 AI 风险治理、映射、测量和管理。"),
    ("NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf", "生成式 AI 风险、测试、来源、事件和生命周期治理。"),
    ("《生成式人工智能服务管理暂行办法》", "https://www.cac.gov.cn/2023-07/13/c_1690898327029107.htm", "生成式 AI 服务、个人信息、内容安全和用户权利。"),
    ("《人工智能生成合成内容标识办法》", "https://www.nrta.gov.cn/art/2025/3/14/art_113_70340.html", "生成合成内容显式/隐式标识和发布声明。"),
    ("GB 45438—2025《网络安全技术 人工智能生成合成内容标识方法》", "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=F32EA2A561F1886CD8D606513512D547", "生成内容标识技术方法。"),
    ("Agent Skills Specification", "https://agentskills.io/specification", "SKILL.md、scripts、references、assets 和渐进式披露。"),
    ("Model Context Protocol Specification 2025-06-18", "https://modelcontextprotocol.io/specification/2025-06-18/basic/index", "JSON-RPC、生命周期、授权、资源、提示和工具。"),
]
ref_table = doc.add_table(rows=1, cols=3)
ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref_table.autofit = False
ref_table.style = "Table Grid"
set_table_borders(ref_table)
for index, (label, width) in enumerate(zip(["参考文件", "链接", "采用重点"], [6.0, 4.0, 7.2])):
    cell = ref_table.rows[0].cells[index]
    set_cell_width(cell, width)
    set_cell_shading(cell, HEADER_FILL)
    set_cell_margins(cell, 90, 70, 90, 70)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(label)
    set_run(run, 7.8, True, INK)
set_repeat_header(ref_table.rows[0])
for name, url, note in references:
    row = ref_table.add_row()
    prevent_row_split(row)
    widths = [6.0, 4.0, 7.2]
    for index, width in enumerate(widths):
        set_cell_width(row.cells[index], width)
        set_cell_shading(row.cells[index], WHITE)
        set_cell_margins(row.cells[index], 45, 65, 45, 65)
    run = row.cells[0].paragraphs[0].add_run(name)
    set_run(run, 7.0, False, INK)
    set_paragraph(row.cells[0].paragraphs[0], 0, 0, 1.05)
    add_hyperlink(row.cells[1].paragraphs[0], "查看官方文件", url)
    set_paragraph(row.cells[1].paragraphs[0], 0, 0, 1.05)
    run = row.cells[2].paragraphs[0].add_run(note)
    set_run(run, 7.0, False, INK)
    set_paragraph(row.cells[2].paragraphs[0], 0, 0, 1.05)
add_note(doc, "版本管理：课程组和平台运营方应至少每学年复核一次模型、协议、案例、标识、安全规则和评测集。若法规、国家标准或行业规范更新，应以现行有效版本为准，并形成变更记录。")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
