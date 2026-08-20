from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "实训资料" / "产品与交付" / "AI通识培训与教学应用平台项目报价单.docx"

BLUE = "000000"
BLUE_DARK = "000000"
TABLE_HEADER = "D9D9D9"
TOTAL_FILL = "EDEDED"
VERY_LIGHT = "F7F7F7"
WHITE = "FFFFFF"
BLUE_SOFT = VERY_LIGHT
BLUE_PALE = WHITE
CYAN_SOFT = VERY_LIGHT
VIOLET_SOFT = VERY_LIGHT
ORANGE_SOFT = VERY_LIGHT
AMBER_SOFT = VERY_LIGHT
GOLD = BLUE
LINE = "BFBFBF"
LINE_DARK = "808080"
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=4, outer_color=LINE_DARK, outer_size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:color"), outer_color if edge in ("top", "left", "bottom", "right") else color)
        node.set(qn("w:sz"), str(outer_size if edge in ("top", "left", "bottom", "right") else size))


def set_paragraph_border(paragraph, edge="bottom", color=BLUE, size=8, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:color"), color)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    page_num = sect_pr.find(qn("w:pgNumType"))
    if page_num is None:
        page_num = OxmlElement("w:pgNumType")
        sect_pr.append(page_num)
    page_num.set(qn("w:start"), str(start))


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def set_run_font(run, size=9, bold=False, color=None, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.2):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(document, text="", size=9, bold=False, color=INK, align=None, before=0, after=3, line=1.35):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    set_paragraph_spacing(p, before=before, after=after, line=line)
    return p


def add_rich_paragraph(document, parts, size=9, before=0, after=3, line=1.35, align=None):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, bold, color in parts:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color or INK)
    set_paragraph_spacing(p, before=before, after=after, line=line)
    return p


def add_heading(document, text, level=1):
    size = 14 if level == 1 else 11
    p = document.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=INK, name="黑体")
    set_paragraph_spacing(p, before=4 if level == 1 else 5, after=6 if level == 1 else 3, line=1.15)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullets(document, items, size=8.8, compact=False):
    for index, item in enumerate(items, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        r = p.add_run(f"（{index}）{item}")
        set_run_font(r, size=size, color=INK)
        set_paragraph_spacing(p, after=1 if compact else 2, line=1.25)


def add_numbered(document, items, size=8.8):
    for index, item in enumerate(items, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        r = p.add_run(f"{index}.  {item}")
        set_run_font(r, size=size, color=INK)
        set_paragraph_spacing(p, after=2, line=1.3)


def add_callout(document, text, fill=BLUE_SOFT, border=BLUE):
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, size=8.7, color=INK)
    set_paragraph_spacing(p, before=2, after=3, line=1.35)


def add_table(document, headers, rows, widths, font_size=8, total_rows=None):
    total_rows = set(total_rows or [])
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, TABLE_HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 95, 80, 95, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(value))
        set_run_font(r, size=font_size, bold=True, color=INK)
        set_paragraph_spacing(p, line=1.1)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        is_total = row_index in total_rows
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, 80, 75, 80, 75)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_total:
                set_cell_shading(cell, TOTAL_FILL)
            else:
                set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if idx == len(values) - 1 and any(ch.isdigit() for ch in str(value)):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=is_total, color=INK)
            set_paragraph_spacing(p, line=1.15)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_break(document):
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)


def add_page_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("AI 通识培训与教学应用平台项目报价单")
    set_run_font(r, size=7.8, color=MUTED)
    set_paragraph_spacing(hp, after=3)
    set_paragraph_border(hp, edge="bottom", color=LINE_DARK, size=4, space=4)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("第 ")
    set_run_font(r, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    r2 = fp.add_run(" 页，共 ")
    set_run_font(r2, size=8, color=MUTED)
    total_fld = OxmlElement("w:fldSimple")
    total_fld.set(qn("w:instr"), "NUMPAGES")
    fp._p.append(total_fld)
    r3 = fp.add_run(" 页")
    set_run_font(r3, size=8, color=MUTED)


doc = Document()
section = doc.sections[0]
configure_page(section)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.4

for style_name in ("List Bullet", "List Number"):
    style = doc.styles[style_name]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(9)

doc.core_properties.title = "AI 通识培训与教学应用平台项目报价单"
doc.core_properties.subject = "100 名教师、80 人峰值 AI 并发、三年混合部署正式报价"
doc.core_properties.author = "AI 通识培训与教学应用平台项目组"

# 1. Cover
add_text(doc, "AI 通识培训与教学应用平台", size=22, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, before=55, after=10, line=1.1)
add_text(doc, "项目报价单", size=24, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=55, line=1.1)
cover_rows = [
    ["报价编号", "QJ-AILOS-20260731-01", "报价日期", "2026 年 7 月 31 日"],
    ["采购单位", "以正式合同及盖章信息为准", "报价单位", "以正式合同及盖章信息为准"],
    ["建设周期", "合同生效后 20 周", "报价有效期", "30 个自然日"],
    ["服务期限", "终验后 3 年，7×8 服务", "适用规模", "100 名教师，峰值 80 人调用 AI"],
    ["含税报价总额", "¥ 2,000,000.00", "人民币大写", "人民币贰佰万元整"],
]
meta = add_table(doc, ["项目", "内容", "项目", "内容"], cover_rows, [2.4, 6.2, 2.4, 6.2], font_size=9)
add_text(doc, "", after=12)
add_callout(doc, "本报价为固定总价控制方案，已包含软件开发、硬件扩容、模型资源、实施、培训、税费和三年运维。税率及发票类型按合同签署时供应商纳税资质执行，含税总价不因税率变化而增加。")

# 2. Summary
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
configure_page(body_section)
add_page_header_footer(body_section)
set_page_number_start(body_section, 1)
add_heading(doc, "一、报价摘要与建设边界")
add_text(doc, "本项目面向 100 名高职教师，建设一套支持 48 课时 AI 通识课程备课授课、教师 AI 能力成长、知识库与 AI 助教、测评评价、工具实训、教学研究和学校教学管理的生产级平台。系统采用“校内业务与数据 + 云端模型弹性调用”的混合部署架构。", size=9.5, color=INK, after=6)
summary_rows = [
    ["正式账号数", "100 个", "使用对象均为教师；管理权限由其中获授权人员承担"],
    ["峰值并发人数", "80 人", "峰值同时调用 AI 的教师人数；应用网关按 100 路能力设计"],
    ["课程建设规模", "48 课时", "每门课程 16 个模块、24 次课，覆盖备课、授课、测评和复盘"],
    ["服务及留存期限", "3 年", "生成内容留存、模型资源池、软件维护和 7×8 技术服务"],
]
add_table(doc, ["项目", "指标", "说明"], summary_rows, [4.0, 3.0, 10.2], font_size=8.5)
add_callout(doc, "总体结论：在甲方已有机房、基础服务器、GPU、存储、防火墙、公网出口和固定 IP 的条件下，200 万元可支撑本项目建设。新增硬件重点用于业务集群、对象存储和独立备份，不重复采购大规模 GPU。")
add_heading(doc, "1. 建设依据", level=2)
add_bullets(doc, [
    "《高职人工智能通识与职业应用》48 课时课程大纲。",
    "100 名教师、共 100 个正式账号，峰值 80 人同时调用 AI；管理权限由其中获授权人员承担。",
    "允许通过公网调用经甲方批准的云端模型，生成内容保存 3 年。",
    "项目预算包含软件、硬件、模型费、实施、培训、税费和三年运维。",
    "服务方式为 7×8，不提供 7×24 驻场或全天候值守。",
], compact=True)
add_heading(doc, "2. 明确边界", level=2)
add_bullets(doc, [
    "正式交付以《AI 通识培训与教学应用平台产品参数说明书》为产品与功能范围依据，并完成后端、数据库、权限、文件存储、模型网关、日志审计、备份和运维体系建设。",
    "本报价不包含独立学生端产品功能及学生直接调用 AI 的账号和模型资源；如需新增，由双方另行确认扩容范围和费用。",
    "不包含旗舰大模型全量私有化部署、模型训练、专属 GPU 集群和第三方等保测评机构费用。",
], compact=True)

# 3. Total quotation
add_page_break(doc)
add_heading(doc, "二、含税报价总表")
quote_rows = [
    ["1", "平台软件产品化", "生产级前后端、教师教学与个人成长双空间、课程备授课、知识库与 AI 助教、测评评价、AI 工具、实训研究、学校教学管理和可信 AI 治理", "80.00", "40.0%"],
    ["2", "系统集成、安全与数据治理", "统一认证接口、模型网关、权限、审计、敏感信息治理、监控告警、备份集成和课程资源初始化", "25.00", "12.5%"],
    ["3", "硬件与备份扩容", "双应用/API 节点、对象存储节点、独立备份节点、10GbE 连接与必要配件，三年质保", "25.00", "12.5%"],
    ["4", "三年模型与多模态资源池", "文本模型、图片、短视频、向量化、文档解析及企业并发资源；通过统一网关计量、路由和限额", "25.00", "12.5%"],
    ["5", "上线实施与教师培训", "部署、初始化、试点、管理员培训、教师应用培训、课程资源导入指导和上线保障", "15.00", "7.5%"],
    ["6", "三年 7×8 运维服务", "故障响应、巡检、备份核验、小版本升级、模型适配、使用报告和年度健康检查", "20.00", "10.0%"],
    ["7", "项目风险与价格波动预留", "用于模型价格、接口变化、必要适配和经双方确认的小范围需求调整，按变更单受控使用", "10.00", "5.0%"],
    ["", "含税总价", "人民币贰佰万元整", "200.00", "100%"],
]
add_table(doc, ["序号", "报价项", "主要内容", "金额（万元）", "占比"], quote_rows, [1.0, 3.6, 8.8, 2.3, 1.5], font_size=7.5, total_rows={7})
add_heading(doc, "报价说明", level=2)
add_numbered(doc, [
    "上述金额均为含税价，包含完成约定建设内容所需的软件、设备、运输、安装、调试、实施、培训和三年服务。",
    "硬件品牌和具体型号在合同技术附件中确定，但性能不得低于本报价单列出的最低配置。",
    "模型资源采用预算池管理，按实际模型、输入输出 Tokens、多媒体次数和厂商结算规则扣减；平台提供按教师、课程、功能和模型的成本统计。",
    "风险预留资金仅用于双方书面确认的项目内变化；未发生部分在结算或续期时按合同约定处理。",
], size=8.5)
add_callout(doc, "本报价不以低价模型长期在售为前提。若模型厂商调整产品或价格，乙方应优先通过同等级替代模型、缓存、批处理和模型路由维持约定容量及核心体验；涉及明显超出资源池的新增使用量，由双方另行确认。", fill=AMBER_SOFT, border="DFA43A")

# 4. Software
add_page_break(doc)
add_heading(doc, "三、平台软件产品化明细")
software_rows = [
    ["1.1", "基础平台与账号权限", "100 个正式账号（均为教师）、对象化任务首页、角色与管理权限、教学/成长双空间、课程与班级上下文、任务中心、消息与个人资料；预留标准统一身份接口", "12.00"],
    ["1.2", "AI 课程备课与教学实施", "48 课时课程模板、16 模块/24 次课编排、情境配置、教案与活动、课件资源、任务发布、版本和引用管理", "22.00"],
    ["1.3", "知识库与 AI 助教", "资料导入、解析、切分、索引、检索测试、课程引用、知识边界、低置信度转人工、助教配置和运营数据", "15.00"],
    ["1.4", "命题、测评、批改与报告", "七类题型、题库、组卷、量规、AI 辅助批改、教师复核、原分与终分、课程报告及资源回流", "12.00"],
    ["1.5", "AI 工具、实训与智能体", "13 类 AI 工具、11 项平台实训任务、18 项教师真实工作任务，以及提示词、多模态、RAG、Agent Skill、MCP、智能体、模型评测、安全测试、四阶段教学研究和成果归档", "11.00"],
    ["1.6", "学校管理与可信 AI 治理", "运行总览、课程班级、学习质量、测评运行、教师培训、AI 用量与成本、成果资产、风险预警、统一证据、模型版本、权限审计、人工确认和导出", "8.00"],
    ["", "平台软件产品化小计", "", "80.00"],
]
add_table(doc, ["序号", "模块", "交付范围", "金额（万元）"], software_rows, [1.0, 4.0, 10.0, 2.2], font_size=7.7, total_rows={6})
add_heading(doc, "主要软件交付物", level=2)
deliver_items = [
    "生产环境 Web 应用和服务端 API", "关系数据库、对象存储与向量索引结构",
    "教师端、学校管理端和课程配置能力", "模型网关、队列、限流和成本台账",
    "部署文档、接口文档和数据字典", "管理员、教师及运维操作手册",
    "测试、压力测试和安全检查报告", "可恢复的安装包、配置和数据库脚本",
]
deliver_rows = [
    [str(i + 1), deliver_items[i], str(i + 2), deliver_items[i + 1]]
    for i in range(0, len(deliver_items), 2)
]
add_table(doc, ["序号", "交付物", "序号", "交付物"], deliver_rows, [1.0, 7.6, 1.0, 7.6], font_size=8.2)
add_text(doc, "", after=2)
add_callout(doc, "数据归属：甲方课程资料、教师输入、生成内容、评分结果、运行日志和导出档案归甲方所有。乙方不得将业务数据用于训练公共模型，云端调用应执行最小必要传输、脱敏和供应商数据保护设置。")

# 5. Hardware
add_page_break(doc)
add_heading(doc, "四、硬件扩容清单与部署架构")
hardware_rows = [
    ["3.1", "应用/API 计算节点", "机架式；≥24 物理核心；≥128GB ECC；2×1.92TB 企业级 SSD RAID1；双 10GbE；冗余电源；三年质保", "2 台", "5.50", "11.00"],
    ["3.2", "对象存储节点", "≥8 盘位；原始容量≥48TB；RAID6/纠删码；可用容量≥10TB；SSD 缓存；10GbE；快照和校验", "1 套", "7.50", "7.50"],
    ["3.3", "独立备份节点", "原始容量≥32TB；可用容量≥20TB；支持增量、全量、保留策略和恢复校验；与生产存储逻辑隔离", "1 套", "5.00", "5.00"],
    ["3.4", "配件与安装材料", "10GbE 光模块/网卡、线缆、磁盘备件、导轨、电源及安装调试材料", "1 批", "1.50", "1.50"],
    ["", "硬件与备份扩容小计", "", "", "", "25.00"],
]
add_table(doc, ["序号", "设备", "最低配置要求", "数量", "单价", "小计"], hardware_rows, [0.9, 2.7, 8.6, 1.2, 1.5, 1.5], font_size=7.2, total_rows={4})
arch_rows = [
    ["校内部署", "账号、权限与课程业务；数据库、对象存储、向量索引；模型网关、用量计费和敏感信息治理；日志审计、监控、备份和恢复；现有 GPU 用于 Embedding、小模型和降级"],
    ["经批准的云端服务", "通用文本生成和复杂推理；视觉理解、图片和短视频生成；高峰并发和弹性容量；向量化、文档解析等可替换能力；只传输完成任务所需的最小数据"],
]
add_table(doc, ["部署位置", "部署内容"], arch_rows, [4.0, 13.2], font_size=8.2)
add_text(doc, "", after=2)
add_text(doc, "硬件采购前由双方完成现有资产盘点。若甲方现有设备达到上述性能，可将对应硬件款经书面变更用于存储扩容、备份、安全或模型资源，但项目含税总价不增加。除非盘点和压测证明必要，本项目不新增大型 GPU 训练/推理集群。", size=8.2, color=MUTED)

# 6. Model resources
add_page_break(doc)
add_heading(doc, "五、三年模型资源与容量指标")
model_rows = [
    ["保守", "800 次", "24 万次", "约 13.20 亿", "少量备课与问答"],
    ["标准基准", "1,500 次", "45 万次", "约 47.25 亿", "完整覆盖 24 次课备授课"],
    ["高强度", "3,000 次", "90 万次", "约 171 亿", "长上下文与高频智能体"],
]
add_table(doc, ["场景", "每人每年调用", "三年总调用", "三年总 Tokens", "使用定位"], model_rows, [2.5, 3.1, 3.0, 3.2, 5.4], font_size=8)
capacity_rows = [
    ["三年文本模型容量", "≥65 亿 Tokens", "标准采购容量目标：基准上预留约 30% 系统提示、重试和工具链空间"],
    ["应用侧在途请求能力", "≥100 路", "满足 80 名教师峰值并发"],
    ["文本服务建议可用配额", "150–200 万 TPM", "复杂智能体采用队列"],
    ["多媒体资源", "10 万图 / 8,000 视频", "三年多媒体资源控制目标"],
]
add_table(doc, ["容量项目", "指标", "说明"], capacity_rows, [4.5, 4.0, 8.7], font_size=8.1)
add_heading(doc, "资源池分配建议", level=2)
resource_rows = [
    ["文本模型资源池", "普通问答、课程生成、复杂推理、命题、批改和报告；高低模型自动路由", "15.00"],
    ["图片、视频与语音资源", "课程插图、图文材料、短视频及必要语音能力", "5.00"],
    ["企业并发与服务保障", "企业配额、接口保障、监控和必要的厂商支持", "3.00"],
    ["资源价格与流量预留", "模型升级、重试、临时峰值及厂商价格变化", "2.00"],
    ["", "三年模型与多模态资源池小计", "25.00"],
]
add_table(doc, ["资源项", "用途", "预算（万元）"], resource_rows, [4.8, 9.8, 2.6], font_size=7.8, total_rows={4})
add_callout(doc, "计量规则：平台应按教师、课程、功能、模型、输入 Tokens、输出 Tokens、图片/视频次数记录用量，并设置个人、课程、月度和项目总额四级预算告警。批量报告和非实时生成优先进入异步队列。")

# 7. Implementation
add_page_break(doc)
add_heading(doc, "六、实施周期与交付成果")
add_text(doc, "项目计划周期为合同生效且甲方提供必要环境、账号和资料后的 20 周。可根据校历安排确定试点课程和正式上线日期。", size=9.2)
timeline_rows = [
    ["第 1–2 周", "调研、资产盘点、需求确认、原型冻结"],
    ["第 3–6 周", "架构、账号权限、数据底座和硬件到货"],
    ["第 7–14 周", "核心功能产品化、模型网关和课程能力"],
    ["第 15–17 周", "集成、资源导入、安全和压力测试"],
    ["第 18–20 周", "试点、培训、问题整改和项目验收"],
]
add_table(doc, ["阶段", "主要工作"], timeline_rows, [3.4, 13.8], font_size=8.3)
delivery_rows = [
    ["1", "生产系统", "部署完成的教师平台、管理端、服务端、数据库、对象存储、向量索引、模型网关和任务队列"],
    ["2", "课程模板", "基于 48 课时大纲形成的 16 模块、24 次课课程结构、资源分类、任务模板、量规和评测框架"],
    ["3", "接口与数据", "接口清单、数据字典、模型接入配置、统一身份预留接口、备份和恢复配置"],
    ["4", "硬件环境", "设备到货、上架、网络配置、存储划分、备份策略和资产清单"],
    ["5", "质量文件", "功能测试、兼容性测试、80 人并发测试、安全检查、备份恢复测试和遗留问题清单"],
    ["6", "培训与文档", "管理员培训、教师培训、运维交接、操作手册、部署手册、应急预案和培训记录"],
]
add_table(doc, ["序号", "成果类别", "交付内容"], delivery_rows, [1.0, 3.5, 12.7], font_size=7.7)
add_heading(doc, "培训安排", level=2)
add_bullets(doc, [
    "平台管理员与运维培训：不少于 2 场，每场不少于 3 小时。",
    "教师备课与授课应用培训：不少于 4 场，每场不少于 3 小时。",
    "课程负责人/骨干教师工作坊：不少于 2 场，每场不少于 3 小时。",
    "提供培训材料、操作录像或可重复使用的图文教程。",
], compact=True)

# 8. Acceptance
add_page_break(doc)
add_heading(doc, "七、项目验收标准")
acceptance_rows = [
    ["1", "账号与权限", "支持 100 个正式账号，使用对象均为教师；管理权限由获授权人员承担；关键操作记录审计", "功能测试"],
    ["2", "课程备授课", "支持 16 模块、24 次课的资源、教案、任务、题库、量规和版本管理", "场景验收"],
    ["3", "AI 能力", "完成文本、知识库、助教、命题、批改、报告及多模态的真实接口调用和结果保存", "接口与场景测试"],
    ["4", "并发能力", "模拟 80 名教师同时使用 AI；平台无崩溃、无数据串扰；请求可流式返回或进入可见队列", "压力测试报告"],
    ["5", "配额治理", "支持模型路由、限流、超时、重试、预算告警和按账号/课程统计", "故障注入测试"],
    ["6", "内容留存", "课程资料和生成内容保存 3 年；支持查询、下载、归档和按权限删除", "数据测试"],
    ["7", "安全与审计", "密钥不下发浏览器；敏感操作留痕；传输加密；关键 AI 结果保留模型、版本和人工复核", "安全检查"],
    ["8", "备份恢复", "每日增量、每周全量、独立备份；抽样完成数据库及对象文件恢复", "恢复演练"],
    ["9", "浏览器适配", "主流桌面浏览器可用；1440px 和 390px 无关键遮挡或横向溢出", "浏览器验收"],
    ["10", "文档培训", "约定文档完整，培训场次完成，管理员能够独立完成账号、课程、模型和配额配置", "资料与实操"],
]
add_table(doc, ["序号", "验收项", "主要标准", "验收方式"], acceptance_rows, [0.9, 3.2, 10.3, 2.8], font_size=7.2)
add_heading(doc, "性能口径", level=2)
add_bullets(doc, [
    "普通业务页面在校内网络正常条件下，95% 请求响应时间不高于 3 秒。",
    "AI 文本任务以流式首字或任务进入队列作为响应；外部模型完整生成时间应可监控、超时和重试。",
    "图片、视频、批量报告和长任务采用异步处理，用户可查看状态、结果和失败原因。",
    "上游厂商实际并发额度不低于合同技术附件，若限流则平台应自动排队并避免请求丢失。",
], compact=True)
add_callout(doc, "终验条件：全部关键功能、80 人并发、数据隔离、备份恢复和培训交接通过；一般问题形成双方确认的整改计划且不影响正式使用。")

# 9. Service & terms
add_page_break(doc)
add_heading(doc, "八、服务等级、付款与商务条款")
add_heading(doc, "1. 三年 7×8 服务", level=2)
service_rows = [
    ["P1 紧急", "系统整体不可用、重大数据风险", "30 分钟内", "4 小时内提供恢复或绕行方案，持续跟踪至恢复"],
    ["P2 高", "核心功能受阻、较大范围用户受影响", "2 小时内", "1 个服务日内提供修复或可用替代方案"],
    ["P3 一般", "局部功能异常，有替代操作", "4 小时内", "3 个服务日内处理或进入最近维护版本"],
    ["P4 咨询", "配置、使用咨询和优化建议", "1 个服务日内", "双方确认的计划内答复或安排"],
]
add_table(doc, ["级别", "典型影响", "首次响应", "处理目标"], service_rows, [2.2, 4.5, 3.0, 7.5], font_size=7.7)
add_text(doc, "7×8 指每日 8 小时、每周 7 天，具体服务时段在合同中确定；法定节假日安排以双方服务日历为准。计划维护应提前通知。", size=8.2, color=MUTED)
add_heading(doc, "2. 建议付款节点", level=2)
payment_rows = [
    ["首付款", "40%", "合同生效，项目启动并提交实施计划", "80.00"],
    ["阶段款", "30%", "硬件到货验收，核心系统完成并进入联调", "60.00"],
    ["初验款", "20%", "试点上线、培训完成并通过初步验收", "40.00"],
    ["终验款", "10%", "整改完成、资料移交并通过最终验收", "20.00"],
]
add_table(doc, ["节点", "比例", "付款条件", "金额（万元）"], payment_rows, [2.5, 2.0, 9.7, 3.0], font_size=8)
add_heading(doc, "3. 不包含项目", level=2)
add_bullets(doc, [
    "学生账号及学生直接调用 AI 产生的扩容；第三方教务、财务等非标准接口的深度改造。",
    "大模型训练、微调、旗舰模型全量私有化部署及新增大型 GPU 集群。",
    "第三方等保测评、密码测评、渗透测试机构、商用 CA 或短信等外部服务费用。",
    "超出资源池的模型、多媒体、搜索、数字人、第三方软件订阅和版权素材费用。",
    "替甲方编写全部校本讲义、录制完整课程视频或承担课程内容的学术与版权终审。",
], compact=True)

# 10. Signature
add_page_break(doc)
add_heading(doc, "九、报价确认")
add_text(doc, "本报价单与后续签署的合同、产品参数说明书、实施计划和验收方案共同构成项目商务与技术依据。如文件之间存在差异，以双方盖章确认的合同及其附件为准。", size=9.5, color=INK, after=6)
add_callout(doc, "有效性说明：本报价自报价日期起 30 个自然日内有效。采购单位名称、报价单位名称、统一社会信用代码、开票信息、收款账户、具体税率、设备品牌型号、服务时段及模型厂商企业配额，应在签约前补充到合同附件并由双方盖章确认。", fill=AMBER_SOFT, border="DFA43A")
sign = doc.add_table(rows=1, cols=2)
sign.alignment = WD_TABLE_ALIGNMENT.CENTER
sign.autofit = False
set_table_borders(sign, color=LINE, size=4, outer_color=LINE_DARK, outer_size=6)
for idx, title in enumerate(("采购单位确认", "报价单位确认")):
    cell = sign.cell(0, idx)
    set_cell_width(cell, 8.6)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 170, 180, 170, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True, color=INK, name="黑体")
    for label in ("单位名称", "授权代表", "联系电话", "日期", "盖章"):
        p = cell.add_paragraph()
        r = p.add_run(f"{label}：____________________________")
        set_run_font(r, size=9, color=MUTED)
        set_paragraph_spacing(p, before=3, after=4)
add_text(doc, "", after=4)
add_heading(doc, "附件清单", level=2)
add_numbered(doc, [
    "产品参数说明书及平台功能范围。",
    "硬件最低技术规格和现有资产盘点表。",
    "模型资源计量、并发指标与费用管理规则。",
    "项目实施计划、培训计划和验收测试用例。",
    "三年运维服务方案及故障响应流程。",
], size=8.8)
add_callout(doc, "建议签约前最后确认：现有 CPU、内存、GPU 型号与显存、存储容量和备份结构；学生是否完全不直接调用平台 AI；采购单位及报价单位的法定名称、发票类型和盖章信息。")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
