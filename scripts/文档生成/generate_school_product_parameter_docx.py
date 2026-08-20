from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "实训资料" / "产品与交付" / "AI通识培训与教学应用平台产品参数说明书.docx"

BLUE = "1F4E79"
HEADER_FILL = "D9D9D9"
LIGHT_FILL = "F7F7F7"
TOTAL_FILL = "EDEDED"
WHITE = "FFFFFF"
LINE = "BFBFBF"
INK = RGBColor(0, 0, 0)
BLUE_RGB = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
FONT = "Microsoft YaHei"

REGISTERED_IDS = []


def set_run(run, size=9, bold=False, color=INK, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_paragraph(paragraph, before=0, after=2, line=1.22):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_paragraph_border(paragraph, color=BLUE, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_heading(doc, value, level=1):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(value)
    if level == 1:
        set_run(run, 16, True, BLUE_RGB)
        set_paragraph(paragraph, 5, 7, 1.1)
        set_paragraph_border(paragraph)
    else:
        set_run(run, 11.3, True, INK)
        set_paragraph(paragraph, 6, 4, 1.1)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_text(doc, value, size=9, bold=False, color=INK, align=None, before=0, after=3, line=1.28):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(value)
    set_run(run, size, bold, color)
    set_paragraph(paragraph, before, after, line)
    return paragraph


def add_bullets(doc, items, size=8.8):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run(run, size, False, INK)
        set_paragraph(paragraph, 0, 1, 1.22)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=LINE, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    tr_pr.append(cannot_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers, rows, widths, font_size=7.35, center_cols=None, total_rows=None, register_ids=False):
    center_cols = center_cols or set()
    total_rows = total_rows or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table)

    header_row = table.rows[0]
    repeat_header(header_row)
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = header_row.cells[index]
        set_cell_width(cell, width)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell, 75, 65, 75, 65)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(paragraph, 0, 0, 1.05)
        run = paragraph.add_run(header)
        set_run(run, font_size + 0.25, True, INK)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if register_ids and values and values[0]:
            REGISTERED_IDS.append(values[0])
        for index, (value, width) in enumerate(zip(values, widths)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_shading(cell, TOTAL_FILL if row_index in total_rows else WHITE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph(paragraph, 0, 0, 1.12)
            run = paragraph.add_run(str(value))
            set_run(run, font_size, row_index in total_rows, INK)
    add_text(doc, "", size=2, after=1)
    return table


def add_parameter_table(doc, rows):
    display_rows = []
    for parameter_id, item, description, value, feature in rows:
        display_item = f"★ {item}" if feature else item
        display_rows.append([parameter_id, display_item, description, value])
    return add_table(
        doc,
        ["编号", "产品参数项", "产品参数与功能说明", "学校应用价值"],
        display_rows,
        [1.15, 3.75, 7.85, 4.45],
        font_size=7.2,
        center_cols={0},
        register_ids=True,
    )


def add_note(doc, value):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 17.2)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, 100, 120, 100, 120)
    run = cell.paragraphs[0].add_run(value)
    set_run(run, 8.5, False, INK)
    set_paragraph(cell.paragraphs[0], 0, 0, 1.25)
    add_text(doc, "", size=2, after=1)


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def set_page_number_start(section, start=1):
    section_pr = section._sectPr
    page_number = section_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_pr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— 第 ")
    set_run(run, 8, False, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页 —")
    set_run(run, 8, False, MUTED)


def add_header_footer(section):
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AI 通识培训与教学应用平台产品参数说明书")
    set_run(run, 7.4, False, MUTED)
    set_paragraph_border(header, color="A6A6A6", size=4)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_hyperlink(paragraph, label, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "14")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([fonts, size, color, underline])
    run.append(run_props)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
cover_section = doc.sections[0]
configure_page(cover_section)
cover_section.top_margin = Cm(1.9)
cover_section.bottom_margin = Cm(1.9)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(9)

add_text(doc, "PRODUCT PARAMETER DESCRIPTION", size=8.2, bold=True, color=BLUE_RGB, align=WD_ALIGN_PARAGRAPH.CENTER, before=35, after=24)
add_text(doc, "AI 通识培训与教学应用平台", size=23, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=9, line=1.05)
add_text(doc, "产品参数说明书", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, line=1.05)
add_text(
    doc,
    "面向高职院校教师 AI 能力建设、AI 通识课程教学与学校教学运行管理",
    size=9.2,
    color=MUTED,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    after=24,
)

cover_rows = [
    ["文档版本", "V3.0", "编制日期", "2026 年 7 月 31 日"],
    ["适用规模", "100 名教师 / 100 个正式账号", "AI 峰值并发", "80 人"],
    ["课程基线", "48 课时 / 16 模块 / 24 次课", "部署模式", "校内业务数据 + 云端模型能力"],
    ["数据留存", "不少于 3 年", "服务周期", "三年 7×8 服务"],
]
add_table(doc, ["项目", "产品口径", "项目", "产品口径"], cover_rows, [2.3, 6.3, 2.3, 6.3], font_size=8.2)
add_note(doc, "本说明书用于向学校系统介绍平台定位、产品特色、功能参数、课程承载能力、部署配置和服务能力。所有功能均以正式产品口径表达，便于学校开展项目论证、方案交流和应用规划。")

body_section = doc.add_section(WD_SECTION.NEW_PAGE)
configure_page(body_section)
add_header_footer(body_section)
set_page_number_start(body_section, 1)

# 一、产品定位与学校价值
add_heading(doc, "一、产品定位与学校价值")
add_heading(doc, "1. 产品定位", level=2)
add_text(
    doc,
    "本平台面向高职院校教师 AI 通识教学和教师 AI 能力发展，将课程备课、教学实施、知识库、AI 助教、测评评价、能力培训、工具实训、教学研究、学校管理和可信 AI 治理组织在统一平台中。平台以教师真实工作任务为入口，以课程与能力证据为数据主线，帮助学校形成可持续运行、可量化评价、可沉淀复用的 AI 教育能力体系。",
    size=9,
    after=5,
)
add_text(
    doc,
    "本期产品面向 100 名教师，共 100 个正式账号；管理权限由获授权人员承担。学校管理端完整呈现课程、班级、学习质量、测评运行、教师培训、AI 应用和风险治理数据。本说明书不设置独立学生端产品章节，也不包含学生直接调用模型的账号与资源。",
    size=9,
    after=6,
)

add_heading(doc, "2. 学校建设需求与产品响应", level=2)
value_rows = [
    ["课程建设分散", "课程结构、资源、题库、量规和版本分别管理，难以形成统一资产。", "以课程知识底座连接备课、授课、助教、命题、评价和复盘。"],
    ["教师应用停留在工具体验", "培训结束后缺少真实工作任务、成果和持续评价。", "以摸底、计划、学习、实训、研究、复测和能力证据形成成长闭环。"],
    ["AI 结果难以复核", "回答、教案、题目和评分缺少来源、版本和人工确认。", "统一记录来源、量规、置信度、证据充分度和人工复核。"],
    ["学校难以掌握运行质量", "课程、教师、测评、模型用量和风险信息分散。", "提供教学运行、培训质量、AI 成本、成果资产和预警管理视图。"],
    ["模型与数据治理压力增加", "多模型、多工具带来密钥、隐私、成本和审计风险。", "通过统一模型网关、权限、内容安全、预算和审计实现集中治理。"],
]
add_table(doc, ["学校关注问题", "主要表现", "本平台的解决方式"], value_rows, [3.5, 6.5, 7.2], font_size=7.8)

add_heading(doc, "3. 核心产品特色", level=2)
feature_rows = [
    ["★", "教师工作与个人成长双空间", "教师在同一身份下分别处理课程教学与个人 AI 学习，两个空间共享成果证据但保持课程、笔记、计划和权限边界。"],
    ["★", "课程全流程一体化", "从标准课程、智能备课、课堂实施、课程资源、AI 助教到测评报告形成连续业务链。"],
    ["★", "知识与证据驱动", "知识库同时服务助教、命题、评分和研究；关键 AI 结果均可回到来源、量规、版本和人工判断。"],
    ["★", "教师真实任务化成长", "将课程学习、自由实训、教学应用、评价复核、分层干预、教学研究和复测纳入统一能力档案。"],
    ["★", "前沿 AI 教学能力", "将上下文工程、Agent Skill、RAG、MCP、Harness、Agent Eval、多模态和智能体安全转化为可教学、可实训的任务。"],
    ["★", "学校教学运行管理", "学校可在统一管理视图中掌握课程运行、学习质量、教师发展、AI 用量、风险预警和成果资产。"],
    ["★", "混合部署与模型可替换", "业务数据和治理能力部署于校内，获准模型通过统一网关调用，支持路由、限额、切换和成本管理。"],
]
feature_display_rows = [[f"★ {name}", description] for _, name, description in feature_rows]
add_table(doc, ["核心产品特色", "特色说明"], feature_display_rows, [5.4, 11.8], font_size=7.9)
add_note(doc, "“★”用于标识平台差异化特色和关键闭环能力，便于学校快速识别产品优势；未标星参数同样属于平台标准产品能力。")

# 二、产品总体架构
add_heading(doc, "二、产品总体架构参数")
architecture_rows = [
    ["A-01", "混合部署架构", "账号、权限、课程业务、数据库、对象存储、向量索引、模型网关、日志和备份部署于校内或学校指定环境；获准的文本、视觉、图片和音视频模型通过云端接口调用。", "兼顾数据可控、模型能力和弹性并发。", "★"],
    ["A-02", "分层产品架构", "产品由访问层、教师业务层、学校管理层、AI 能力层、数据与知识层、集成层、安全治理层和基础设施层组成。", "便于分层建设、运维和持续扩展。", ""],
    ["A-03", "统一模型网关", "所有模型请求经服务端统一网关完成路由、限流、超时、重试、熔断、预算、成本统计和模型切换，浏览器端不保存模型密钥。", "集中管理模型质量、成本与安全。", "★"],
    ["A-04", "异步任务中心", "图片、视频、知识库解析、批量报告和批量评价采用异步任务，提供排队、进度、取消、失败原因、重试和完成通知。", "高峰期任务可见、可控、不丢失。", ""],
    ["A-05", "统一数据与知识底座", "课程、资源、题目、量规、任务、作品、报告、能力证据和研究成果使用统一对象模型并保留来源与版本。", "减少信息孤岛，形成学校数字资产。", "★"],
    ["A-06", "开放集成", "提供标准 REST API，并预留统一身份、教务、课程平台、对象存储和 MCP 等开放协议的集成能力。", "便于接入学校现有信息化环境。", ""],
    ["A-07", "运行可观测性", "集中监控业务接口、模型调用、任务队列、Tokens、费用、失败率、低置信度、人工接管、存储和备份状态。", "支持主动发现问题和持续优化。", ""],
    ["A-08", "终端与无障碍适配", "适配主流 Chromium 内核桌面浏览器，关键页面兼容 1440px 桌面和 390px 移动视口，支持键盘操作、清晰焦点和语义标题。", "满足办公室、教室和移动查看场景。", ""],
]
add_parameter_table(doc, architecture_rows)

# 三、教师工作空间
add_heading(doc, "三、教师工作空间与任务入口参数")
workspace_rows = [
    ["W-01", "教师账号规模", "标准配置支持 100 名教师、共 100 个正式账号，并可按学校、院系、专业、课程和教师组进行组织。", "形成清晰的教师应用范围和组织结构。", ""],
    ["W-02", "授权管理权限", "可从正式账号中授权平台管理、教学管理和运维管理权限，并按菜单、对象、数据和操作进行配置。", "兼顾管理效率和最小权限。", ""],
    ["W-03", "对象化任务首页", "登录后优先展示课程、班级、课次、测评、报告、实训和研究等真实对象，显示状态、时间、指标和唯一主操作。", "教师从待办任务直接进入工作，降低学习成本。", "★"],
    ["W-04", "课程与班级上下文", "教师选择的课程和班级可在备课、教学、学习促进、测评和诊断之间继承，切换对象时清理不兼容的测评与报告选择。", "减少重复选择和对象错配。", "★"],
    ["W-05", "教师双空间", "教师身份包含“教学工作”和“我的成长”两个独立空间；个人学习课程、笔记、实训和能力报告不覆盖教学班级上下文。", "同时支撑教师教学与个人发展。", "★"],
    ["W-06", "六类教师任务", "教学工作空间提供智能教学、学习促进、测评管理、教学诊断、AI 研究和自由实训六类入口。", "按照教师真实任务组织平台功能。", ""],
    ["W-07", "状态化任务分流", "测评和课程任务根据待配置、待发布、进行中、待评价、待复核和已完成等状态进入对应操作页面。", "让教师直接处理当前阶段任务。", ""],
    ["W-08", "多课程隔离", "课程、班级、任务、资源、测评、报告和补学状态使用课程主键隔离；不同课程之间不串用数据和上下文。", "支持学校并行运行多门课程。", "★"],
]
add_parameter_table(doc, workspace_rows)

# 四、课程建设与智能教学
add_heading(doc, "四、课程建设与智能教学参数")
teaching_rows = [
    ["T-01", "课程结构配置", "课程可配置为模块、课次和课时三级结构；每次课包含目标、知识点、流程、资源、活动、任务、题目、量规和课后安排。", "把课程设计转化为可运行的教学单元。", "★"],
    ["T-02", "标准课程模板", "提供 48 课时、16 个模块、24 次课的 AI 通识课程结构，并支持学校按专业和校本要求复制、调整和形成版本。", "缩短课程从规划到开课的准备周期。", "★"],
    ["T-03", "智能备课流程", "备课支持标准方案选择、教学情境配置、方案生成、方案对比、教师修改、采纳发布和版本留痕。", "保留教师主导权，提高备课效率。", "★"],
    ["T-04", "教学情境配置", "可设置教学对象、基础水平、人数、授课方式、课时、目标难度和教学重点，并将其作为方案适配依据。", "同一课程可适配不同院系和班级。", ""],
    ["T-05", "方案差异说明", "展示标准方案与适配方案在目标、难度、节奏、活动、资源、题目和进度方面的差异及依据。", "让教师理解并判断 AI 调整结果。", "★"],
    ["T-06", "教师采纳与版本", "AI 方案先保存为草稿，教师修改确认后发布；系统同时保留 AI 原稿、教师修改、发布版本、时间和操作人。", "确保课程方案可复核、可回退。", "★"],
    ["T-07", "四类备课成果", "采纳方案后形成情境化教案、课堂活动包、配套题目集和教学进度方案四类结构化成果。", "一次备课形成完整的教学实施材料。", ""],
    ["T-08", "教学实施工作台", "展示可衡量目标、适配知识点、课堂流程、关联资源、分层题目和准备度检查，支持步骤时长、资源和题目调整。", "帮助教师按课次组织课堂。", ""],
    ["T-09", "教学任务发布", "教学任务发布前提供目标、资源、题目、截止时间和对象检查，发布后进入课程运行记录。", "降低遗漏并形成教学过程数据。", ""],
    ["T-10", "课程资源中心", "资源分为标准资源、AI 生成资源和校本资源，记录来源、版本、状态、权限、标签、知识点和引用关系。", "形成可管理、可复用的课程资源库。", "★"],
    ["T-11", "资源导入与校验", "支持课程大纲、讲义、案例、题目、量规、任务单和术语表导入，并显示解析、重复、覆盖和异常信息。", "帮助学校快速初始化校本课程资料。", ""],
    ["T-12", "课程版本管理", "课程、模块、课次、方案和资源支持草稿与发布版本隔离、版本对比、历史引用和回退。", "保障课程长期迭代的稳定性。", ""],
    ["T-13", "多模态教学资源", "支持图文、流程图、海报、音频、短视频和数字人脚本等资源生成，记录提示、模型、版本、人工修改、授权和 AI 标识。", "丰富课程呈现并落实版权和标识管理。", ""],
    ["T-14", "教学成果导出", "教案、活动、进度、题目、量规和资源清单可按权限导出，用于教研、归档和线下教学。", "便于成果复用和学校资产沉淀。", ""],
]
add_parameter_table(doc, teaching_rows)

# 五、知识库与 AI 助教
add_heading(doc, "五、课程知识库与 AI 助教参数")
knowledge_rows = [
    ["K-01", "课程知识库", "将课程大纲、讲义、案例、题目、量规和制度文件组织为可检索、可引用的统一课程知识源。", "为助教、命题、评价和研究提供可信依据。", "★"],
    ["K-02", "文档处理", "提供文档解析、分段、去重、元数据、向量化和索引能力，保留文件、章节、页码和更新时间。", "让知识片段可追踪、可维护。", ""],
    ["K-03", "检索配置", "支持关键词、向量和混合检索，可配置召回数量、相似度阈值、重排和知识范围。", "适配不同课程的问答和生成任务。", ""],
    ["K-04", "知识质量检测", "提供覆盖率、重复率、失效内容、解析异常、检索命中和测试集结果。", "帮助课程团队持续提升知识质量。", "★"],
    ["K-05", "课程知识图谱", "将知识点、先修关系、资源、题目、任务和掌握证据形成可视化关联，并支持教师调整。", "直观呈现课程结构和教学覆盖。", ""],
    ["K-06", "AI 助教配置", "配置流程包含知识绑定、教学人格、呈现形态、规则边界、测试和发布；知识库为基础，知识图谱作为增强。", "快速形成面向课程的专属 AI 助教。", "★"],
    ["K-07", "有引用回答", "助教回答展示课程来源、章节、知识点和置信度，引用可返回原始知识片段。", "便于教师核验回答依据。", "★"],
    ["K-08", "知识边界", "超出知识范围、证据冲突或置信度不足时，助教明确提示边界并转入教师确认。", "减少无依据的确定性回答。", "★"],
    ["K-09", "发布与开放范围", "助教发布前检查知识、版本、测试集、开放范围、内容安全和人工接管规则，发布后保留入口和版本信息。", "实现受控上线和按课程使用。", ""],
    ["K-10", "助教运营分析", "统计解决率、知识命中率、高频问题、低置信度、越界问题、人工接管、调用量和成本。", "支持课程团队优化知识和助教规则。", ""],
]
add_parameter_table(doc, knowledge_rows)

# 六、学习促进与教学诊断
add_heading(doc, "六、学习促进与教学诊断参数")
learning_rows = [
    ["L-01", "多周期学习观察", "教师可按当前课次、本周和本单元查看参与、完成、进度、作品和能力变化，过程数据与正式成绩分开呈现。", "支持形成性评价和及时干预。", ""],
    ["L-02", "班级态势", "集中展示参与度、完成率、进度分布、目标达成和可行动信号，并可查看对应课程与课次。", "帮助教师快速识别需要关注的范围。", ""],
    ["L-03", "多源学习证据", "按目标聚合任务、课堂练习、作品迭代、测评过程、助教互动和教师评价，显示来源、充分度和更新时间。", "以证据支撑教学判断。", "★"],
    ["L-04", "诊断结果区分", "明确区分“尚未掌握”和“证据不足”；低置信度、证据冲突或证据不足结果进入教师复核。", "避免将数据不足误判为能力不足。", "★"],
    ["L-05", "动态分组", "根据目标差距和证据将教学支持组织为补强、巩固和拓展三类动态分组。", "让分层教学具有明确依据。", ""],
    ["L-06", "教师可编辑干预", "教师可调整干预目标、资源、题目、难度、时长、截止时间和接收对象后发布。", "确保 AI 建议符合真实教学安排。", "★"],
    ["L-07", "干预结果回流", "任务完成、作品变化和教师确认结果回流班级态势、目标诊断和下一轮教学建议。", "形成发现问题、实施干预、观察变化的闭环。", "★"],
    ["L-08", "教学诊断报告", "提供班级、课程、知识点和匿名个体层级的成绩、达成、误区、作品变化、复核和改进建议。", "把过程数据转化为可执行的教学改进。", ""],
]
add_parameter_table(doc, learning_rows)

# 七、测评评价
add_heading(doc, "七、测评、评价与报告参数")
assessment_rows = [
    ["E-01", "测评场景", "支持摸底测、随堂测、单元测、正式考试、实训和综合项目等多类测评。", "覆盖教师培训和课程教学的不同评价阶段。", ""],
    ["E-02", "七类题型", "题库支持单选、多选、判断、填空、简答、案例分析和创作实操七类题型。", "兼顾知识、判断、方法和真实成果评价。", ""],
    ["E-03", "课程题库", "统一管理 AI 生成、教师创建和历史题目，记录来源、知识点、题型、难度、答案、解析、量规、版本和使用效果。", "形成可复用、可持续改进的测评资产。", ""],
    ["E-04", "智能命题", "根据课程目标、知识点、题型、难度和认知层级生成题目、答案、解析、来源和评分量规，教师确认后入库。", "提高命题效率并保留专业审核。", "★"],
    ["E-05", "组卷蓝图", "先设置总分、知识覆盖、难度、题型、认知层级和评分主体，再自动匹配题目并检查相似题与量规。", "确保试卷结构与课程目标一致。", "★"],
    ["E-06", "状态化测评管理", "待发布、进行中、待评价和已完成测评分别进入组卷发布、进度监控、AI 辅助评价和教师复核。", "减少教师查找和切换成本。", ""],
    ["E-07", "多主体评分", "综合项目支持机器客观评分、AI 作品评价和教师人工评分，标准比例为 30/50/20，并允许课程级配置。", "兼顾效率、开放作品评价和教师判断。", "★"],
    ["E-08", "AI 评价证据", "AI 作品评价展示任务符合度、专业准确性初检、方法与工具、迭代证据、内容安全、量规条目和置信度。", "让评价过程透明、可复核。", "★"],
    ["E-09", "教师最终复核", "教师可修改 AI 分、人工分和评语，最终分实时重算，同时保留 AI 原值、修改理由和最终结果。", "落实教师对评价结果的最终责任。", "★"],
    ["E-10", "质量检查", "提供总分、覆盖、难度、相似题、量规、异常评分、低置信度和人工复核率检查。", "提升测评和评价质量。", ""],
    ["E-11", "报告体系", "提供提交、成绩、达成、复核、难度、分布、知识掌握、共性误区和改进建议，并可下钻到原始证据。", "支持课程复盘和教学改进。", ""],
    ["E-12", "过程性评价", "能力评价综合课程学习、知识测评、实训作品、教学应用、研究成果和复测，不以登录次数或单次考试替代。", "形成更真实的教师能力评价。", "★"],
]
add_parameter_table(doc, assessment_rows)

# 八、教师成长
add_heading(doc, "八、教师 AI 能力成长参数")
growth_rows = [
    ["G-01", "成长全流程", "形成摸底测评、能力报告、培训计划、课程学习、自主学习、AI 实训、教学应用、研究成果、成长报告和复测的连续闭环。", "将一次培训转化为持续能力发展。", "★"],
    ["G-02", "5+3 能力模型", "五个计分维度为 AI 基础认知、提示词与多模态、知识库与智能体、教学融合、研究创新；设置事实核验、数据版权和人工责任三项通关门槛。", "兼顾能力水平和可信应用底线。", "★"],
    ["G-03", "L1—L4 等级", "支持 L1 入门、L2 实践、L3 创新和 L4 引领四级能力等级，等级由测评、任务、作品和应用证据共同支撑。", "便于学校制定分层培养目标。", ""],
    ["G-04", "摸底测评", "摸底包含理论、情境和实操证据，可记录选择、提示词、截图、文件、测试结果和人工确认。", "识别教师真实起点而非自我评价。", ""],
    ["G-05", "个性化培训计划", "根据能力报告从已配置课程生成计划，教师可调整学习顺序、每周投入和选修内容。", "提高培训与个人短板的匹配度。", ""],
    ["G-06", "课程与自主学习", "同时提供 AI 推荐路径和教师自主路径，记录资源进度、收藏、笔记、任务和继续学习位置。", "兼顾学校组织和教师自主发展。", ""],
    ["G-07", "七类能力证据", "课程学习、自由实训、教学实施、评价复核、分层干预、教学研究和复测等成果进入个人证据台账。", "把培训成果与真实工作表现连接起来。", "★"],
    ["G-08", "证据下钻", "能力报告可查看证据来源、任务、作品、时间、量规、得分、教师确认和版本。", "让能力结论可解释、可复核。", ""],
    ["G-09", "教师真实工作任务", "平台配置不少于 18 项教师真实工作任务，覆盖备课资源、命题评价、知识库智能体和教学研究四条任务链。", "以可交付成果促进能力迁移。", "★"],
    ["G-10", "任务成果模板", "任务包含目标、输入、操作、成果、量规、前置、难度、时长、风险边界和人工确认字段。", "统一培训组织和成果质量要求。", ""],
    ["G-11", "成长报告", "报告展示综合等级、五维能力、任务达成、代表成果、证据充分度、优势、短板和下一步建议。", "支持教师自我发展和学校培养决策。", ""],
    ["G-12", "延迟复测", "支持首次摸底、过程评价、结业复测和 30/90 天延迟复测，展示能力变化和工作迁移证据。", "观察培训效果是否稳定迁移。", ""],
    ["G-13", "成果归档", "教师可将课程、实训、教学和研究成果按权限归档为个人或学校资产，保留版本和复用记录。", "沉淀可持续使用的校本成果。", ""],
    ["G-14", "个人数据边界", "个人答案、笔记、计划和研究草稿默认仅本人可见；学校管理视图使用授权范围内的统计和成果信息。", "保护教师个人发展空间。", "★"],
]
add_parameter_table(doc, growth_rows)

# 九、AI 工具、实训与研究
add_heading(doc, "九、AI 工具、实训与教学研究参数")
research_rows = [
    ["R-01", "共享工具入口", "研究与自由实训使用同一工具目录，支持分类筛选、工具详情、起步任务、难度、时长、权限和启动记录。", "避免工具孤岛，统一教师使用体验。", "★"],
    ["R-02", "深度研究助手", "支持研究问题、检索范围、证据矩阵、来源核验、冲突证据、过程记录和人工结论。", "提升资料研究的规范性和可追溯性。", ""],
    ["R-03", "文献研读与证据库", "支持文献摘要、观点、证据、引用、主题标签和研究问题关联。", "形成可复用的教学研究证据。", ""],
    ["R-04", "研究数据分析室", "支持匿名数据导入、描述统计、交叉分析、图表、异常说明和人工解释。", "帮助教师用数据验证教学问题。", ""],
    ["R-05", "提示词实验室", "记录任务、上下文、提示版本、输出、评价、修改和最佳实践。", "培养可复用的上下文与提示设计能力。", ""],
    ["R-06", "多模态内容工坊", "提供图片、图文、语音和多模态内容生成，记录模型、提示、版本、授权、人工修改和标识。", "支持课程资源和成果表达。", ""],
    ["R-07", "文生视频工坊", "支持教学短视频脚本、镜头、字幕、生成任务、版本和披露信息。", "降低微课和演示资源制作门槛。", ""],
    ["R-08", "知识库 RAG 实训", "提供资料导入、检索配置、测试集、引用核验和质量报告。", "让教师掌握可信知识增强方法。", ""],
    ["R-09", "低代码应用实验", "支持自然语言或可视化方式组合模型、知识、工具、表单、条件和结果页面。", "帮助教师形成轻量教学应用。", ""],
    ["R-10", "智能体搭建", "可配置目标、角色、指令、知识、Skill、工具、状态、规划、权限、护栏和人工接管。", "把智能体能力转化为可控教学应用。", "★"],
    ["R-11", "MCP 工具连接", "支持在受控环境理解和配置资源、提示与工具，展示输入输出、权限、调用轨迹和失败信息。", "培养开放工具连接和权限意识。", ""],
    ["R-12", "多模态识别", "支持图片、文档、表格和界面内容识别，并保留来源、识别结果和人工校对。", "服务资源整理、分析和无障碍应用。", ""],
    ["R-13", "Vibe Coding 工作台", "支持自然语言生成轻量网页或工具原型，保留需求、验收标准、依赖、测试、修改和安全检查。", "帮助非开发教师快速验证教学创意。", ""],
    ["R-14", "模型评测场", "使用固定评测集对不同模型或版本进行重复测试，记录质量、引用、耗时、成本、失败和人工接管。", "为学校选模和模型切换提供依据。", "★"],
    ["R-15", "十一项平台实训任务", "提供创建知识库、检索优化、生成图谱、配置智能体、数字人、发布助教、智能出题、设计实训、智能组卷、AI 评价和学情分析任务链。", "把平台核心能力转化为可完成的实训。", ""],
    ["R-16", "成果回流", "实训记录工具、任务、时长、得分、作品、版本和人工确认，并回流教师能力报告。", "让每次实训形成可评价成果。", "★"],
    ["R-17", "四阶段教学研究", "研究项目按问题设计、证据汇集、分析验证、成果复核四阶段推进，保留来源、参数和人工判断。", "支持教师围绕真实教学问题开展研究。", "★"],
    ["R-18", "科研智能体", "支持绑定研究项目和资料，配置角色、工具、测试集、版本和权限，并关联运行记录与研究成果。", "提高研究过程组织和工具复用效率。", ""],
]
add_parameter_table(doc, research_rows)

# 十、学校教学管理
add_heading(doc, "十、学校教学运行管理参数")
management_rows = [
    ["M-01", "运行总览", "集中展示课程、教师、班级、任务、测评、报告、AI 应用、预警和服务健康等指标，并支持组织范围筛选。", "为学校提供统一运行视图。", "★"],
    ["M-02", "课程与班级管理", "查看课程版本、班级状态、教师安排、课次进度、资源、任务和异常，支持按院系、专业和课程筛选。", "掌握课程运行和教学组织情况。", ""],
    ["M-03", "学习质量", "统计目标达成、证据充分度、任务完成、作品迭代、知识点趋势和干预效果；个体明细默认匿名。", "从结果和过程共同判断教学质量。", ""],
    ["M-04", "测评运行", "统计待发布、进行中、待评价、待复核和已完成测评，以及提交、成绩、复核、异常和改进情况。", "及时发现测评组织和评价问题。", ""],
    ["M-05", "教师培训管理", "展示培训参与、计划完成、任务达成、能力分布、成果类型、复测变化和支持需求。", "支持教师分层培养和资源配置。", "★"],
    ["M-06", "AI 应用与成本", "展示 Tokens、费用、模型、成功率、响应、引用命中、低置信度、人工接管、内容安全和工具调用指标。", "控制模型成本并监测应用质量。", "★"],
    ["M-07", "成果资产管理", "统一归档课程模板、提示词、量规、知识库、Skill、智能体、评测集和研究成果，支持审核、复用和版本管理。", "形成学校可持续复用的 AI 教育资产。", ""],
    ["M-08", "风险预警", "预警记录对象范围、证据、责任人、截止时间、状态和处置结果，支持交办、跟踪、解决和指标更新。", "把教学和 AI 风险转化为管理行动。", "★"],
    ["M-09", "教师干预闭环", "管理关注可转化为教师任务，教师完成教学干预和人工确认后，结果回流管理指标和审计记录。", "形成学校管理与教师改进的闭环。", "★"],
    ["M-10", "统计与导出", "按学期、课程、院系和项目导出运行、培训、成本、风险、成果和改进数据，并执行权限和匿名检查。", "支持项目汇报、归档和持续改进。", ""],
]
add_parameter_table(doc, management_rows)

# 十一、前沿 AI 教学设计
add_heading(doc, "十一、前沿 AI 教学设计参数")
frontier_rows = [
    ["F-01", "上下文工程", "将任务目标、角色、受众、输入资料、约束、示例、输出格式和评价标准组织为可复用上下文模板。", "帮助教师稳定获得符合教学要求的结果。", "★"],
    ["F-02", "结构化输出", "支持表格、JSON 或 Schema 约束，检查字段、类型、必填、异常和失败回退。", "便于将 AI 结果进入课程、题库和管理流程。", ""],
    ["F-03", "Agent Skill", "支持按照说明、脚本、参考资料和资产组织可复用技能，并记录适用范围和版本。", "沉淀教师和学校的可复用工作方法。", "★"],
    ["F-04", "RAG 教学设计", "将资料准备、分段、检索、引用、测试集、失败分析和优化组织为完整实训流程。", "让知识增强从概念学习进入真实应用。", "★"],
    ["F-05", "智能体工作流", "使用目标、知识、Skill、工具、状态、规则、护栏和人工确认节点组织可视工作流。", "培养教师面向复杂任务的流程设计能力。", ""],
    ["F-06", "MCP 开放连接", "在受控环境中呈现资源、提示和工具的连接方式、权限范围、输入输出和调用轨迹。", "理解标准化工具连接和安全边界。", ""],
    ["F-07", "状态与记忆", "支持任务状态、短期上下文、受控记忆、断点恢复、幂等和清理规则。", "提升长任务和多步骤应用的稳定性。", ""],
    ["F-08", "Harness", "把运行环境、上下文、知识、工具、规则、日志、反馈和恢复机制组织为可查看的运行配置。", "帮助教师理解智能体可靠运行所需条件。", "★"],
    ["F-09", "Agent Eval", "通过固定评测集重复运行，记录成功率、引用、耗时、成本、工具调用、人工接管和失败类别。", "让智能体改进有统一评价依据。", "★"],
    ["F-10", "多模态教学设计", "将文本、图片、语音、视频、识别和无障碍要求组织为教学资源与成果任务。", "扩展教师内容设计和表达方式。", ""],
    ["F-11", "计算机使用智能体", "在仿真或沙箱页面中展示观察、计划、操作、确认、失败和恢复，高风险操作由人工确认。", "理解智能体执行操作的能力与风险。", ""],
    ["F-12", "深度研究", "支持问题分解、检索范围、证据矩阵、来源核验、冲突证据、过程记录和人工结论。", "提升教师信息研究和证据判断能力。", ""],
    ["F-13", "Vibe Coding", "以自然语言生成轻量应用原型，并结合需求、验收标准、测试、依赖和安全检查进行迭代。", "支持教师快速形成教学工具原型。", ""],
]
add_parameter_table(doc, frontier_rows)

# 十二、可信治理
add_heading(doc, "十二、可信 AI、数据安全与治理参数")
security_rows = [
    ["S-01", "统一 AI 证据", "关键 AI 结果记录来源、知识点、输入摘要、模型与提示版本、量规、结论、置信度、证据充分度、权限、匿名状态、人工复核和修改时间。", "让回答、方案、题目和评价可解释、可复核。", "★"],
    ["S-02", "证据状态", "提供草稿、待复核、已确认、已发布、已退回和证据不足等状态，并记录状态变化。", "明确 AI 结果所处责任阶段。", ""],
    ["S-03", "低置信度治理", "低置信度、无引用、证据冲突或证据不足结果显式提示并进入人工处理。", "避免不可靠结果直接影响教学和管理。", "★"],
    ["S-04", "人工控制", "课程发布、评分修改、外部发送、工具写入、成果公开和高风险内容均由授权人员确认。", "保持教师和学校的最终决策权。", "★"],
    ["S-05", "个人信息保护", "按最小必要原则处理教师、课程、班级和研究数据，支持查阅、更正、导出、删除和分级权限。", "降低个人信息使用风险。", ""],
    ["S-06", "数据归属与使用", "学校课程资料、教师输入、生成内容、评分结果、日志和导出档案归学校所有，不用于训练公共模型。", "保障学校数据权益。", "★"],
    ["S-07", "密钥与模型安全", "模型、数据库、存储和工具密钥由服务端管理，支持轮换、隔离、白名单和泄露处置。", "避免密钥暴露和越权调用。", ""],
    ["S-08", "内容安全", "对违法违规、暴力、歧视、隐私、版权、虚假信息和高风险专业建议进行检测、提示、拦截或转人工。", "降低教学内容和生成内容风险。", ""],
    ["S-09", "版权与生成标识", "记录来源、授权、引用和人工修改，对生成或合成内容提供显式和隐式标识能力。", "便于规范使用和成果披露。", ""],
    ["S-10", "工具权限分级", "工具按只读或写入、可逆或不可逆、低中高风险分级，调用前展示权限和影响。", "控制智能体工具使用边界。", ""],
    ["S-11", "审计日志", "记录登录、权限、课程发布、模型调用、工具调用、评分修改、数据导出、删除、配置变化和异常事件。", "支持安全追溯和责任确认。", "★"],
    ["S-12", "事件处置", "风险事件支持登记、分级、责任人、处理过程、恢复、复盘和规则更新，并关联模型、课程和影响范围。", "形成持续改进的风险治理机制。", ""],
]
add_parameter_table(doc, security_rows)

# 十三、容量部署硬件
add_heading(doc, "十三、容量、部署与硬件配置参数")
add_text(doc, "以下配置面向 100 名教师、峰值 80 人同时调用 AI 的标准应用场景，并结合学校已有机房、服务器、GPU、存储、防火墙、公网出口和固定 IP 进行配置。", size=8.8, color=MUTED, after=5)
capacity_rows = [
    ["C-01", "账号容量", "100 名教师、共 100 个正式账号；管理权限由获授权人员承担。", "本期标准授权规模"],
    ["C-02", "AI 峰值并发", "支持 80 名教师同时发起 AI 任务，应用网关按不低于 100 路在途请求配置。", "为交互请求和后台任务保留余量"],
    ["C-03", "文本服务配额", "文本模型可用配额建议为 150—200 万 TPM，复杂智能体任务通过队列控制扇出。", "兼顾峰值体验与上游配额"],
    ["C-04", "文本模型资源", "三年文本模型资源池不低于 65 亿 Tokens，包含系统提示、重试、工具链和增长余量。", "覆盖完整备授课和教师成长场景"],
    ["C-05", "多媒体资源", "三年支持不少于 10 万次图片生成、8,000 次短视频生成或学校确认的等值资源池。", "支撑课程资源和多模态实训"],
    ["C-06", "业务响应", "校内网络正常条件下，普通业务页面 95% 请求响应时间不高于 3 秒。", "保障常规业务操作体验"],
    ["C-07", "AI 任务响应", "文本任务支持流式返回；长任务进入可见队列并显示状态、取消、超时、重试和失败原因。", "外部模型波动时任务仍可管理"],
    ["C-08", "内容留存", "课程资料、生成内容、成果、评分、报告和关键审计数据保存不少于 3 年。", "满足项目周期内查询与归档"],
    ["C-09", "对象存储", "可用对象存储容量不低于 10TB，支持版本、校验、生命周期、归档和权限下载。", "承载课程资源和生成成果"],
    ["C-10", "备份存储", "独立备份可用容量不低于 20TB，并与生产存储逻辑隔离。", "支持三年数据保护和恢复"],
    ["C-11", "网络与安全", "使用固定 IP 和 HTTPS，对外模型调用经防火墙策略控制；服务器和存储采用 10GbE 连接。", "满足混合部署和安全访问"],
    ["C-12", "监控与预算", "提供个人、课程、月度和项目总额四级预算提醒，并监控模型、任务、存储、备份和服务健康。", "支持学校持续控制资源和风险"],
]
add_table(doc, ["编号", "配置项", "标准产品参数", "适用说明"], capacity_rows, [1.2, 3.3, 9.1, 3.6], font_size=7.4, center_cols={0}, register_ids=True)

add_heading(doc, "标准硬件推荐配置", level=2)
hardware_rows = [
    ["H-01", "应用/API 计算节点", "机架式；≥24 物理核心；≥128GB ECC；2×1.92TB 企业级 SSD RAID1；双 10GbE；冗余电源；三年质保。", "2 台"],
    ["H-02", "对象存储节点", "≥8 盘位；原始容量≥48TB；RAID6 或纠删码；可用容量≥10TB；SSD 缓存；10GbE；快照和校验。", "1 套"],
    ["H-03", "独立备份节点", "原始容量≥32TB；可用容量≥20TB；支持增量、全量、保留策略和恢复校验；与生产存储逻辑隔离。", "1 套"],
    ["H-04", "配件与安装材料", "10GbE 光模块或网卡、线缆、磁盘备件、导轨、电源及安装调试材料。", "1 批"],
]
add_table(doc, ["编号", "设备", "标准推荐配置", "数量"], hardware_rows, [1.2, 3.5, 10.9, 1.6], font_size=7.4, center_cols={0, 3}, register_ids=True)
add_note(doc, "硬件配置前结合学校现有资产完成盘点。现有设备达到同等能力时，可优先复用并将新增资源用于存储、备份、安全或模型服务扩容。除非盘点和压力测试证明必要，本场景不要求新增大型 GPU 训练或推理集群。")

# 十四、课程映射
add_heading(doc, "十四、48 课时课程内容与平台能力映射")
add_text(doc, "课程内容通过目标、知识点、资源、任务、工具、成果、量规和风险边界进入平台，实现课程教学、教师实训和能力评价的一体化组织。", size=8.8, color=MUTED, after=5)
course_rows = [
    ["CM-01", "AI 基础、职业变革与人机分工", "4", "课程导学、课前诊断、深度研究与模型评测", "AI 应用辨识图、岗位任务人机分工表"],
    ["CM-02", "数据、机器学习与模型评价", "4", "研究数据分析、模型评测、随堂测与单元测", "数据质量诊断表、模型选择说明"],
    ["CM-03", "生成式 AI 原理与可信使用", "4", "知识资源、模型评测、风险情境与成长报告", "AI 输出核验记录、可信使用清单"],
    ["CM-04", "上下文工程", "2", "提示词实验室、AI 助教与成果版本", "岗位任务上下文 v1/v2"],
    ["CM-05", "提示词与结构化输出", "4", "提示词实验室、数据分析与模型评测", "三轮提示记录、结构化输出模板"],
    ["CM-06", "Agent Skill 设计", "2", "提示词实验室、Vibe Coding 与成果模板", "入门级 SKILL.md、试用反馈"],
    ["CM-07", "多模态内容生成", "2", "多模态工坊、识别、版本与版权检查", "多模态作品、无障碍检查记录"],
    ["CM-08", "语音与数字人", "2", "文生视频、多模态工坊与数字人形态", "语音或数字人方案、字幕与披露记录"],
    ["CM-09", "计算机使用智能体", "2", "仿真页面、智能体搭建与工具连接", "受控操作轨迹、人工确认清单"],
    ["CM-10", "课程知识库与 RAG", "4", "课程资源、知识库、AI 助教与检索测试", "来源清单、6—9 题检索测试报告"],
    ["CM-11", "工具调用与 MCP", "2", "MCP 工具连接、低代码应用与权限矩阵", "工具说明书、输入输出样例、权限表"],
    ["CM-12", "状态、记忆与任务恢复", "2", "智能体搭建、低代码应用与运行轨迹", "状态流转图、断点恢复和幂等测试"],
    ["CM-13", "智能体设计与低代码开发", "4", "智能体搭建、低代码、MCP 与护栏", "智能体原型、权限风险矩阵、人机交接"],
    ["CM-14", "Harness 与运行可观测性", "2", "Vibe Coding、模型评测、轨迹与日志", "Harness 画布、运行轨迹标注"],
    ["CM-15", "智能体评测与安全", "4", "模型评测、MCP、红队测试与评价证据", "不少于 8 例评测集、红队报告、原型 v2"],
    ["CM-16", "专业智能体综合项目", "4", "成果版本、模型评测、成长报告与答辩", "AI 职业应用成果包、治理清单"],
    ["", "合计", "48", "16 个模块 / 24 次课", "完整 AI 职业应用成果包"],
]
for row in course_rows[:-1]:
    REGISTERED_IDS.append(row[0])
add_table(doc, ["编号", "模块主题", "课时", "平台能力映射", "核心成果"], course_rows, [1.25, 4.15, 0.95, 6.0, 4.85], font_size=7.15, center_cols={0, 2}, total_rows={16})

# 十五、产品交付与服务
add_heading(doc, "十五、产品交付、培训与服务参数")
delivery_rows = [
    ["D-01", "项目设计", "提供需求确认、产品配置、总体架构、部署拓扑、数据结构、接口、安全、备份和项目计划。", "形成学校认可的建设蓝图"],
    ["D-02", "生产系统", "提供可运行的 Web 应用、服务端、数据库、对象存储、向量索引、模型网关、任务队列和管理端。", "具备正式运行所需的完整系统"],
    ["D-03", "课程初始化", "完成 48 课时课程结构、资源分类、任务模板、量规、测评框架和示例资源配置。", "快速建立课程运行基础"],
    ["D-04", "内容服务边界", "标准服务包含课程结构、模板、量规、评测框架和示例资源；完整校本讲义、完整课程视频及学术版权终审由学校课程团队负责。", "明确平台产品与校本内容职责"],
    ["D-05", "数据与接口资料", "提供数据字典、接口清单、模型适配配置、统一身份预留接口、导入导出和备份恢复配置。", "支持学校集成和后续运维"],
    ["D-06", "资源导入指导", "提供大纲、讲义、案例、题目、量规和制度资料的导入模板、校验规则和异常处理指导。", "帮助学校完成校本资源初始化"],
    ["D-07", "质量测试", "覆盖教师教学、成长、测评、研究、学校管理、课程隔离、权限、证据、预警、80 人并发、安全和备份恢复。", "验证核心场景的稳定性与完整性"],
    ["D-08", "培训服务", "平台管理与运维培训不少于 2 场、每场不少于 3 小时；教师应用培训不少于 4 场、每场不少于 3 小时；骨干教师工作坊不少于 2 场、每场不少于 3 小时。", "保障管理人员和教师能够独立使用"],
    ["D-09", "试点运行", "选择代表课程完成不少于 1 个完整备课、教学、测评、报告和改进闭环。", "在正式推广前验证学校应用方式"],
    ["D-10", "建设周期", "标准建设周期为合同生效且学校提供必要环境、账号和资料后的 20 周。", "便于结合校历安排部署和试点"],
]
add_table(doc, ["编号", "服务项", "服务参数", "学校获得"], delivery_rows, [1.2, 3.5, 9.0, 3.5], font_size=7.35, center_cols={0}, register_ids=True)

add_heading(doc, "三年运维服务", level=2)
service_rows = [
    ["O-01", "服务方式", "提供三年 7×8 技术服务，具体服务时段和节假日安排在服务计划中确认。", "持续技术支持和运行保障"],
    ["O-02", "P1 紧急事件", "首次响应不高于 30 分钟，4 小时内提供恢复或绕行方案，并持续跟踪至恢复。", "重大故障快速恢复"],
    ["O-03", "P2 高级事件", "首次响应不高于 2 小时，1 个服务日内提供修复或可用替代方案。", "核心功能问题及时处理"],
    ["O-04", "P3 一般事件", "首次响应不高于 4 小时，3 个服务日内处理或进入最近维护版本。", "一般问题有明确处理计划"],
    ["O-05", "P4 咨询事件", "首次响应不高于 1 个服务日，按确认计划答复或安排。", "配置和使用问题获得支持"],
    ["O-06", "日常服务内容", "包括故障响应、巡检、备份核验、小版本升级、模型适配、使用报告和年度健康检查。", "持续优化平台运行质量"],
    ["O-07", "升级与回滚", "支持灰度发布、版本回滚、数据库迁移、模型版本切换和维护通知，升级保留历史课程与证据。", "降低升级对教学运行的影响"],
]
add_table(doc, ["编号", "服务项", "服务参数", "学校获得"], service_rows, [1.2, 3.5, 9.0, 3.5], font_size=7.35, center_cols={0}, register_ids=True)

# 十六、覆盖索引
add_heading(doc, "十六、产品能力覆盖索引")
coverage_rows = [
    ["总体架构与平台底座", "A-01—A-08", "混合部署、模型网关、异步任务、数据知识底座、开放集成、监控与终端适配"],
    ["教师工作空间", "W-01—W-08", "账号权限、对象任务首页、上下文、双空间、状态分流和多课程隔离"],
    ["课程建设与智能教学", "T-01—T-14", "课程模板、智能备课、教学实施、资源、版本、多模态与成果导出"],
    ["课程知识库与 AI 助教", "K-01—K-10", "文档处理、检索、质量、图谱、助教配置、引用、边界和运营"],
    ["学习促进与教学诊断", "L-01—L-08", "班级态势、多源证据、诊断、分组、干预、结果回流和教学报告"],
    ["测评评价与报告", "E-01—E-12", "题型、题库、命题、组卷、评价证据、教师复核和过程性评价"],
    ["教师能力成长", "G-01—G-14", "5+3 模型、L1—L4、摸底、计划、学习、七类证据、任务、报告和复测"],
    ["AI 工具、实训与研究", "R-01—R-18", "13 类工具、11 项平台任务、成果回流、四阶段研究和科研智能体"],
    ["学校教学运行管理", "M-01—M-10", "运行、课程班级、学习质量、测评、教师培训、成本、资产、预警和导出"],
    ["前沿 AI 教学设计", "F-01—F-13", "上下文、Skill、RAG、MCP、Harness、Agent Eval、多模态、研究和 Vibe Coding"],
    ["可信 AI 与数据治理", "S-01—S-12", "证据、低置信度、人工控制、隐私、数据、内容安全、工具权限和审计"],
    ["容量、硬件与服务", "C-01—C-12 / H-01—H-04 / D-01—D-10 / O-01—O-07", "100 名教师、80 人并发、模型资源、存储备份、交付培训和三年服务"],
    ["48 课时课程映射", "CM-01—CM-16", "16 个模块、24 次课及对应平台能力和核心成果"],
]
add_table(doc, ["能力领域", "参数范围", "覆盖内容"], coverage_rows, [4.3, 4.2, 8.7], font_size=7.4)
add_note(doc, "本索引用于学校快速定位产品能力。平台将教师教学、个人成长、学校管理和可信 AI 治理连接为连续运行体系，避免功能分散为互不关联的工具入口。")

# 十七、规范性依据
add_heading(doc, "十七、规范性依据与版本管理")
references = [
    ("JY/T 0646—2022《教师数字素养》", "https://www.moe.gov.cn/srcsite/A16/s3342/202302/W020230214594527529113.pdf", "教师数字意识、知识技能、应用、责任与专业发展。"),
    ("UNESCO AI Competency Framework for Teachers (2024)", "https://www.unesco.org/en/articles/ai-competency-framework-teachers", "人本、AI 伦理、基础应用、AI 教学法、专业学习和能力进阶。"),
    ("NIST AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework", "可信 AI 风险识别、测量和管理。"),
    ("NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf", "生成式 AI 风险、测试、来源、事件和生命周期治理。"),
    ("《生成式人工智能服务管理暂行办法》", "https://www.cac.gov.cn/2023-07/13/c_1690898327029107.htm", "生成式 AI 服务、个人信息、内容安全和用户权益。"),
    ("《人工智能生成合成内容标识办法》", "https://www.nrta.gov.cn/art/2025/3/14/art_113_70340.html", "生成合成内容显式与隐式标识和发布声明。"),
    ("GB 45438—2025《人工智能生成合成内容标识方法》", "https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=F32EA2A561F1886CD8D606513512D547", "生成内容标识技术方法。"),
    ("Agent Skills Specification", "https://agentskills.io/specification", "SKILL.md、脚本、参考资料、资产和渐进式披露。"),
    ("Model Context Protocol Specification 2025-06-18", "https://modelcontextprotocol.io/specification/2025-06-18/basic/index", "生命周期、授权、资源、提示和工具。"),
]
ref_table = doc.add_table(rows=1, cols=3)
ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref_table.autofit = False
set_table_borders(ref_table)
for index, (label, width) in enumerate(zip(["参考文件", "官方链接", "产品采用重点"], [6.0, 3.8, 7.4])):
    cell = ref_table.rows[0].cells[index]
    set_cell_width(cell, width)
    set_cell_shading(cell, HEADER_FILL)
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(label)
    set_run(run, 7.5, True, INK)
repeat_header(ref_table.rows[0])
for name, url, note in references:
    row = ref_table.add_row()
    prevent_row_split(row)
    for index, width in enumerate([6.0, 3.8, 7.4]):
        set_cell_width(row.cells[index], width)
        set_cell_margins(row.cells[index], 45, 65, 45, 65)
    run = row.cells[0].paragraphs[0].add_run(name)
    set_run(run, 6.9, False, INK)
    set_paragraph(row.cells[0].paragraphs[0], 0, 0, 1.03)
    add_hyperlink(row.cells[1].paragraphs[0], "官方资料", url)
    set_paragraph(row.cells[1].paragraphs[0], 0, 0, 1.03)
    run = row.cells[2].paragraphs[0].add_run(note)
    set_run(run, 6.9, False, INK)
    set_paragraph(row.cells[2].paragraphs[0], 0, 0, 1.03)
add_text(
    doc,
    "版本管理：按学年复核模型、协议、案例、生成标识、安全规则和评测集；法规或标准更新时同步更新平台配置并保留版本记录。",
    size=7.2,
    color=MUTED,
    before=2,
    after=0,
    line=1.0,
)

# 生成前一致性检查
if len(REGISTERED_IDS) != len(set(REGISTERED_IDS)):
    duplicates = sorted({value for value in REGISTERED_IDS if REGISTERED_IDS.count(value) > 1})
    raise ValueError(f"参数编号重复: {duplicates}")

course_hours = sum(int(row[2]) for row in course_rows[:-1])
if course_hours != 48:
    raise ValueError(f"课程课时合计错误: {course_hours}")

full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
full_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

required_phrases = [
    "100 名教师",
    "80 名教师",
    "100 路",
    "150—200 万 TPM",
    "65 亿 Tokens",
    "10TB",
    "20TB",
    "20 周",
    "三年 7×8",
    "48 课时",
    "16 个模块",
    "24 次课",
    "Agent Skill",
    "MCP",
    "Harness",
    "Agent Eval",
    "13 类工具",
    "18 项教师真实工作任务",
]
missing = [phrase for phrase in required_phrases if phrase not in full_text]
if missing:
    raise ValueError(f"缺少关键产品口径: {missing}")

forbidden_phrases = ["启境", "Demo", "甲方", "乙方", "供应商应逐项响应", "响应规则", "验收方式"]
present_forbidden = [phrase for phrase in forbidden_phrases if phrase in full_text]
if present_forbidden:
    raise ValueError(f"出现不符合产品说明语气的内容: {present_forbidden}")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
print(f"registered_ids={len(REGISTERED_IDS)} stars={full_text.count('★')} course_hours={course_hours}")
